#!/usr/bin/env python3
from __future__ import annotations

import sys
import argparse
import json
import os
import re
import shutil
import subprocess
import tempfile
import unicodedata
from collections import OrderedDict
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from urllib.parse import urljoin

sys.dont_write_bytecode = True

try:
    import requests
except Exception:  # pragma: no cover
    requests = None

try:
    from bs4 import BeautifulSoup, NavigableString, Tag
except Exception:  # pragma: no cover
    BeautifulSoup = None
    NavigableString = None
    Tag = None

try:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except Exception:  # pragma: no cover
    OxmlElement = None
    qn = None

from markdown_docx_renderer import render_markdown_to_docx

USER_AGENT = "Mozilla/5.0"
PATENT_BASE_URL = "https://patents.google.com/patent/{publication_number}"
PATENT_TEXT_TABLE_NAME = "patent_text_records.xlsx"
PATENT_TEXT_TABLE_SHEET = "patent_texts"
PATENT_TEXT_CELL_LIMIT = 32000
PATENT_TEXT_TABLE_COLUMNS = [
    "publication_number",
    "source_office",
    "title",
    "source_url",
    "final_url",
    "status_code",
    "crawl_status",
    "docx_path",
    "pdf_path",
    "abstract",
    "claims",
    "description",
    "abstract_length",
    "claims_length",
    "description_length",
    "abstract_truncated",
    "claims_truncated",
    "description_truncated",
    "warnings",
]
IMAGE_EXT_RE = re.compile(r"\.(png|jpe?g|webp|tif|tiff)(\?|$)", re.IGNORECASE)
PDF_EXT_RE = re.compile(r"\.pdf(\?|$)", re.IGNORECASE)
FIGURE_RE = re.compile(
    r"(?:\u56fe|\u5716|\u56f3|\ub3c4|fig(?:s|ures?)?\.?)\s*([\d]{1,3}[A-Za-z]?)",
    re.IGNORECASE,
)
UNICODE_HYPHENS = "\u2010\u2011\u2012\u2013\u2014\u2212"
INLINE_MISSING_PLACEHOLDER = "\u3010\u516c\u5f0f/\u5bf9\u8c61\u56fe\u7f3a\u5931\u3011"
BLOCK_MISSING_PLACEHOLDER = "\u3010\u6b64\u5904\u516c\u5f0f/\u5bf9\u8c61\u56fe\u6e90\u9875\u672a\u63d0\u4f9b\u53ef\u4e0b\u8f7d\u56fe\u7247\u3011"
INLINE_DOWNLOAD_FAILED_PLACEHOLDER = "\u3010\u6b64\u5904\u516c\u5f0f/\u5bf9\u8c61\u56fe\u6293\u53d6\u5931\u8d25\u3011"
DRAWING_MISSING_PLACEHOLDER = "\u3010\u9644\u56fe\u7f3a\u5931\u3011"
CONTENT_MISSING_PLACEHOLDER = "\u3010\u5185\u5bb9\u7f3a\u5931\u3011"
DESCRIPTION_METADATA_MARKERS = (
    "global patent litigation dataset",
    "darts-ip",
    "creative commons attribution",
)
DESCRIPTION_BODY_MARKERS = (
    "\u6280\u672f\u9886\u57df",
    "\u80cc\u666f\u6280\u672f",
    "\u53d1\u660e\u5185\u5bb9",
    "\u5b9e\u7528\u65b0\u578b\u5185\u5bb9",
    "\u5177\u4f53\u5b9e\u65bd\u65b9\u5f0f",
    "\u9644\u56fe\u8bf4\u660e",
)
DRAWING_HEADING_MARKERS = (
    "brief description of the drawings",
    "brief description of drawings",
    "brief description of the several views of the drawings",
    "\u9644\u56fe\u8bf4\u660e",
    "\u56f3\u9762\u306e\u7c21\u5358\u306a\u8aac\u660e",
    "\ub3c4\uba74\uc758 \uac04\ub2e8\ud55c \uc124\uba85",
)
DRAWING_SECTION_END_MARKERS = (
    "detailed description",
    "detailed description of the preferred embodiments",
    "description of embodiments",
    "description of embodiments of the invention",
    "best mode",
    "reference signs list",
    "\u5177\u4f53\u5b9e\u65bd\u65b9\u5f0f",
    "\u5b9e\u65bd\u4f8b",
    "\u7b26\u53f7\u306e\u8aac\u660e",
    "\uc2e4\uc2dc\uc608",
)
OFFICE_RE = re.compile(r"^[A-Za-z]+")


def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _error(msg: str, code: int = 1) -> None:
    print(json.dumps({"ok": False, "error": msg}, ensure_ascii=False, indent=2))
    raise SystemExit(code)


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8-sig")


def _extract_mermaid_from_markdown(text: str) -> str | None:
    fence = re.search(r"```mermaid\s*(.*?)```", text, re.DOTALL | re.IGNORECASE)
    if fence:
        return fence.group(1).strip()
    return None


def _looks_like_mermaid(text: str) -> bool:
    head = "\n".join(text.strip().splitlines()[:5]).lower()
    return any(
        token in head
        for token in (
            "graph ",
            "flowchart",
            "sequencediagram",
            "classdiagram",
            "statediagram",
            "erdiagram",
            "gantt",
            "pie",
            "mindmap",
        )
    )


def _detect_type(path: Path, content: str, forced: str) -> str:
    if forced != "auto":
        return forced

    suffix = path.suffix.lower()
    if suffix in {".mmd", ".mermaid"}:
        return "mermaid"
    if suffix in {".json"}:
        return "json"
    if suffix in {".md", ".markdown"}:
        if _extract_mermaid_from_markdown(content):
            return "mermaid"
        return "markdown"

    if _extract_mermaid_from_markdown(content) or _looks_like_mermaid(content):
        return "mermaid"

    try:
        json.loads(content)
        return "json"
    except Exception:
        return "markdown"


def _strip_markdown_inline(text: str) -> str:
    out = text
    out = re.sub(r"!\[[^\]]*\]\([^\)]*\)", "[image]", out)
    out = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", out)
    out = re.sub(r"`([^`]+)`", r"\1", out)
    out = re.sub(r"\*\*([^*]+)\*\*", r"\1", out)
    out = re.sub(r"__([^_]+)__", r"\1", out)
    out = re.sub(r"\*([^*]+)\*", r"\1", out)
    out = re.sub(r"_([^_]+)_", r"\1", out)
    out = re.sub(r"~~([^~]+)~~", r"\1", out)
    return out


def configure_docx_styles(doc: Any) -> None:
    if qn is None or OxmlElement is None:
        return

    def apply_font(style_name: str) -> None:
        style = doc.styles[style_name]
        style.font.name = "KaiTi"
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        rfonts.set(qn("w:ascii"), "KaiTi")
        rfonts.set(qn("w:hAnsi"), "KaiTi")
        rfonts.set(qn("w:cs"), "KaiTi")
        rfonts.set(qn("w:eastAsia"), "楷体")

    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"):
        apply_font(style_name)


def _markdown_to_docx(text: str, output: Path) -> None:
    try:
        render_markdown_to_docx(text, output, configure_document=configure_docx_styles)
    except RuntimeError as exc:
        _error(str(exc))


def _wrap_cjk_line(text: str, width: int = 54) -> list[str]:
    if not text:
        return [""]
    buf: list[str] = []
    cur = ""
    for ch in text:
        cur += ch
        if len(cur) >= width:
            buf.append(cur)
            cur = ""
    if cur:
        buf.append(cur)
    return buf or [""]


def _markdown_to_pdf(text: str, output: Path) -> None:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.pdfgen import canvas
    except Exception as exc:
        _error(f"missing_dependency: reportlab ({exc})")

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    c = canvas.Canvas(str(output), pagesize=A4)
    width, height = A4
    left = 40
    top = height - 40
    y = top
    line_h = 16

    in_code = False

    def new_page() -> None:
        nonlocal y
        c.showPage()
        c.setFont("STSong-Light", 11)
        y = top

    c.setFont("STSong-Light", 11)
    for raw in text.splitlines():
        line = raw.rstrip("\n")
        if line.strip().startswith("```"):
            in_code = not in_code
            continue

        if not in_code:
            line = re.sub(r"^#{1,6}\s+", "", line)
            line = re.sub(r"^[-*]\s+", "* ", line)
            line = re.sub(r"^\d+\.\s+", "", line)
            line = _strip_markdown_inline(line)

        wrapped = _wrap_cjk_line(line)
        for part in wrapped:
            if y <= 40:
                new_page()
            c.drawString(left, y, part)
            y -= line_h
        if line.strip() == "":
            y -= 4

    c.save()


def _ordered_union_keys(rows: list[dict[str, Any]]) -> list[str]:
    keys: OrderedDict[str, None] = OrderedDict()
    for row in rows:
        for key in row.keys():
            if key not in keys:
                keys[key] = None
    return list(keys.keys())


def _json_to_xlsx(data: Any, output: Path) -> None:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font
    except Exception as exc:
        _error(f"missing_dependency: openpyxl ({exc})")

    wb = Workbook()
    ws = wb.active
    ws.title = "data"

    if isinstance(data, list):
        if all(isinstance(i, dict) for i in data):
            rows = data
            headers = _ordered_union_keys(rows)
            ws.append(headers)
            for item in rows:
                ws.append([item.get(h, "") for h in headers])
        else:
            ws.append(["value"])
            for item in data:
                ws.append([json.dumps(item, ensure_ascii=False)])
    elif isinstance(data, dict):
        list_key = next(
            (
                k
                for k, v in data.items()
                if isinstance(v, list) and v and all(isinstance(i, dict) for i in v)
            ),
            None,
        )
        if list_key:
            rows = data[list_key]
            headers = _ordered_union_keys(rows)
            ws.append(headers)
            for item in rows:
                ws.append([item.get(h, "") for h in headers])
        else:
            ws.append(["key", "value"])
            for k, v in data.items():
                ws.append([k, json.dumps(v, ensure_ascii=False) if isinstance(v, (dict, list)) else v])
    else:
        ws.append(["value"])
        ws.append([json.dumps(data, ensure_ascii=False)])

    if ws.max_row >= 1:
        for col in range(1, ws.max_column + 1):
            cell = ws.cell(row=1, column=col)
            cell.font = Font(bold=True)

    for row in ws.iter_rows():
        for cell in row:
            if cell.value is None:
                continue
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for col in ws.columns:
        max_len = 0
        col_letter = col[0].column_letter
        for cell in col:
            val = "" if cell.value is None else str(cell.value)
            max_len = max(max_len, len(val))
        ws.column_dimensions[col_letter].width = min(max(12, max_len + 2), 60)

    wb.save(output)


def _render_mermaid_to_png(source_path: Path, content: str, output: Path) -> None:
    mmdc = shutil.which("mmdc")
    if not mmdc:
        _error("missing_dependency: mmdc (install @mermaid-js/mermaid-cli)")

    mermaid_text = content
    if source_path.suffix.lower() not in {".mmd", ".mermaid"}:
        extracted = _extract_mermaid_from_markdown(content)
        if extracted:
            mermaid_text = extracted

    tmp_input: Path | None = None
    try:
        if source_path.suffix.lower() in {".mmd", ".mermaid"}:
            input_file = source_path
        else:
            fd, tmp = tempfile.mkstemp(prefix="patent-mermaid-", suffix=".mmd")
            os.close(fd)
            Path(tmp).write_text(mermaid_text, encoding="utf-8")
            tmp_input = Path(tmp)
            input_file = tmp_input

        proc = subprocess.run(
            [mmdc, "-i", str(input_file), "-o", str(output), "-b", "transparent"],
            check=False,
            capture_output=True,
            text=True,
        )
        if proc.returncode != 0:
            err = (proc.stderr or proc.stdout or "mmdc_failed").strip()
            _error(f"mermaid_render_failed: {err}")
    finally:
        if tmp_input and tmp_input.exists():
            tmp_input.unlink(missing_ok=True)


def run_convert_mode(args: argparse.Namespace) -> dict[str, Any]:
    if not args.input:
        _error("input_required_for_convert_mode")

    src = Path(args.input).expanduser().resolve()
    if not src.exists():
        _error(f"input_not_found: {src}")

    content = _read_text(src)
    detected = _detect_type(src, content, args.type)

    out_dir = Path(args.output_dir).expanduser().resolve() if args.output_dir else src.parent
    out_dir.mkdir(parents=True, exist_ok=True)
    base = args.base_name or src.stem

    outputs: list[str] = []
    if detected == "markdown":
        docx_path = out_dir / f"{base}.docx"
        pdf_path = out_dir / f"{base}.pdf"
        _markdown_to_docx(content, docx_path)
        _markdown_to_pdf(content, pdf_path)
        outputs.extend([str(docx_path), str(pdf_path)])
    elif detected == "json":
        try:
            data = json.loads(content)
        except Exception as exc:
            _error(f"json_parse_failed: {exc}")
        xlsx_path = out_dir / f"{base}.xlsx"
        _json_to_xlsx(data, xlsx_path)
        outputs.append(str(xlsx_path))
    elif detected == "mermaid":
        png_path = out_dir / f"{base}.png"
        _render_mermaid_to_png(src, content, png_path)
        outputs.append(str(png_path))
    else:
        _error(f"unsupported_type: {detected}")

    return {
        "ok": True,
        "mode": "convert",
        "input": str(src),
        "detected_type": detected,
        "outputs": outputs,
    }


def load_payload(args: argparse.Namespace) -> dict[str, Any]:
    payload: dict[str, Any] = {}
    if args.answers:
        payload.update(json.loads(args.answers))
    if args.answers_file:
        payload.update(json.loads(Path(args.answers_file).read_text(encoding="utf-8-sig")))
    if args.context:
        raw = json.loads(args.context)
        if isinstance(raw, dict):
            payload.update(raw)
    if args.context_file:
        payload.update(json.loads(Path(args.context_file).read_text(encoding="utf-8-sig")))
    return payload


def flatten_context(payload: dict[str, Any]) -> dict[str, Any]:
    context: dict[str, Any] = {}

    required_inputs = payload.get("required_inputs")
    if isinstance(required_inputs, dict):
        context.update(required_inputs)

    runtime_ctx = payload.get("context_for_prompt_runtime")
    if isinstance(runtime_ctx, dict):
        context.update(runtime_ctx)

    for key, value in payload.items():
        if key in {
            "skill_name",
            "timestamp",
            "required_inputs",
            "stage_plan",
            "validation_errors",
            "ready",
            "parsed_artifacts",
            "context_for_prompt_runtime",
            "warnings",
        }:
            continue
        context[key] = value

    return context


def as_str(value: Any) -> str:
    return "" if value is None else str(value).strip()


def as_bool(value: Any, default: bool) -> bool:
    if isinstance(value, bool):
        return value
    token = as_str(value).lower()
    if token in {"1", "true", "yes", "y", "on"}:
        return True
    if token in {"0", "false", "no", "n", "off"}:
        return False
    return default


def as_int(value: Any, default: int) -> int:
    if isinstance(value, int):
        return value
    token = as_str(value)
    if not token:
        return default
    try:
        return int(token)
    except Exception:
        return default


def as_list(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, list):
        return [str(v).strip() for v in value if str(v).strip()]
    text = as_str(value)
    if not text:
        return []
    if text.startswith("[") and text.endswith("]"):
        try:
            parsed = json.loads(text)
            if isinstance(parsed, list):
                return [str(v).strip() for v in parsed if str(v).strip()]
        except Exception:
            pass
    return [p.strip() for p in re.split(r"[\n,;]+", text) if p.strip()]


def unique_keep_order(items: list[str]) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for item in items:
        if item in seen:
            continue
        seen.add(item)
        out.append(item)
    return out


def build_warning_text(warnings: list[str]) -> str:
    return "; ".join(unique_keep_order([as_str(item) for item in warnings if as_str(item)]))


def truncate_table_text(label: str, text: Any, warnings: list[str]) -> tuple[str, int, bool]:
    normalized = "" if text is None else str(text)
    original_length = len(normalized)
    truncated = original_length > PATENT_TEXT_CELL_LIMIT
    if truncated:
        normalized = normalized[:PATENT_TEXT_CELL_LIMIT]
        warnings.append(f"{label}_truncated:{original_length}")
    return normalized, original_length, truncated


def load_json_file(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def build_patent_text_table_row(result: dict[str, Any]) -> dict[str, Any]:
    publication_number = as_str(result.get("publication_number"))
    row_warnings = unique_keep_order([as_str(item) for item in (result.get("warnings") or []) if as_str(item)])
    parsed_payload: dict[str, Any] = {}

    parsed_json_path = as_str(result.get("parsed_json_path"))
    if parsed_json_path:
        try:
            parsed_payload = load_json_file(Path(parsed_json_path))
        except Exception as exc:
            row_warnings.append(f"parsed_json_read_failed:{exc}")

    sections = parsed_payload.get("sections")
    if not isinstance(sections, dict):
        sections = {}

    abstract, abstract_length, abstract_truncated = truncate_table_text(
        "abstract",
        sections.get("abstract", ""),
        row_warnings,
    )
    claims, claims_length, claims_truncated = truncate_table_text(
        "claims",
        sections.get("claims", ""),
        row_warnings,
    )
    description, description_length, description_truncated = truncate_table_text(
        "description",
        sections.get("description", ""),
        row_warnings,
    )

    return {
        "publication_number": publication_number,
        "source_office": as_str(parsed_payload.get("source_office")) or detect_source_office(publication_number),
        "title": as_str(parsed_payload.get("title")),
        "source_url": as_str(parsed_payload.get("source_url") or result.get("source_url")),
        "final_url": as_str(parsed_payload.get("final_url") or result.get("final_url")),
        "status_code": parsed_payload.get("status_code", ""),
        "crawl_status": "success" if result.get("ok") else "failed",
        "docx_path": as_str(result.get("docx_path")),
        "pdf_path": as_str(result.get("pdf_path")),
        "abstract": abstract,
        "claims": claims,
        "description": description,
        "abstract_length": abstract_length,
        "claims_length": claims_length,
        "description_length": description_length,
        "abstract_truncated": abstract_truncated,
        "claims_truncated": claims_truncated,
        "description_truncated": description_truncated,
        "warnings": build_warning_text(row_warnings),
    }


def write_patent_text_table(rows: list[dict[str, Any]], output: Path) -> None:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font
    except Exception as exc:
        _error(f"missing_dependency: openpyxl ({exc})")

    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = PATENT_TEXT_TABLE_SHEET
    worksheet.append(PATENT_TEXT_TABLE_COLUMNS)

    for row in rows:
        worksheet.append([row.get(column, "") for column in PATENT_TEXT_TABLE_COLUMNS])

    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    body_alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)

    for column_index in range(1, worksheet.max_column + 1):
        cell = worksheet.cell(row=1, column=column_index)
        cell.font = Font(bold=True)
        cell.alignment = header_alignment

    for row in worksheet.iter_rows(min_row=2):
        for cell in row:
            if cell.value is None:
                continue
            cell.alignment = body_alignment

    preferred_widths = {
        "publication_number": 22,
        "source_office": 12,
        "title": 30,
        "source_url": 36,
        "final_url": 36,
        "status_code": 12,
        "crawl_status": 12,
        "docx_path": 36,
        "pdf_path": 36,
        "abstract": 60,
        "claims": 60,
        "description": 60,
        "abstract_length": 16,
        "claims_length": 16,
        "description_length": 18,
        "abstract_truncated": 18,
        "claims_truncated": 18,
        "description_truncated": 20,
        "warnings": 48,
    }

    for column_index, column_name in enumerate(PATENT_TEXT_TABLE_COLUMNS, 1):
        worksheet.column_dimensions[worksheet.cell(row=1, column=column_index).column_letter].width = preferred_widths.get(
            column_name,
            24,
        )

    output.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(output)


def export_patent_text_table(results: list[dict[str, Any]], output_root: Path) -> tuple[str, int]:
    rows = [build_patent_text_table_row(result) for result in results]
    output_path = output_root / PATENT_TEXT_TABLE_NAME
    write_patent_text_table(rows, output_path)
    return str(output_path), len(rows)


def detect_source_office(publication_number: str) -> str:
    match = OFFICE_RE.match(publication_number or "")
    return match.group(0).upper() if match else "UNKNOWN"


def is_drawing_heading_text(text: str) -> bool:
    normalized = normalize_inline_text(text)
    if not normalized:
        return False
    lower = normalized.lower()
    return any(marker in lower for marker in DRAWING_HEADING_MARKERS)


def is_drawing_section_end_text(text: str) -> bool:
    normalized = normalize_inline_text(text)
    if not normalized:
        return False
    lower = normalized.lower()
    return any(marker in lower for marker in DRAWING_SECTION_END_MARKERS)


def looks_like_heading_text(text: str) -> bool:
    normalized = normalize_inline_text(text)
    if not normalized or len(normalized) > 180:
        return False
    if is_drawing_heading_text(normalized) or is_drawing_section_end_text(normalized):
        return True
    if any(marker in normalized for marker in DESCRIPTION_BODY_MARKERS):
        return True
    letters = [ch for ch in normalized if ch.isalpha()]
    if letters and sum(1 for ch in letters if ch.isupper()) >= max(3, int(len(letters) * 0.7)):
        return True
    return False


def is_heading_block(section_key: str, node: Any, text: str, tokens: list[dict[str, Any]]) -> bool:
    if not text:
        return False
    if getattr(node, "name", "").lower() == "heading":
        return True
    if section_key != "description":
        return False
    if any(token.get("type") != "text" for token in tokens):
        return False
    if FIGURE_RE.search(text):
        return False
    return looks_like_heading_text(text)


def iter_section_block_nodes(section_node: Any, section_key: str) -> list[Any]:
    if section_node is None:
        return []

    if section_key == "claims":
        top_level: list[Any] = []
        seen: set[int] = set()
        for selector in (
            "li.claim > div.claim",
            "li.claim-dependent > div.claim",
            "claim",
            "div.claim",
        ):
            for node in section_node.select(selector):
                marker = id(node)
                if marker in seen:
                    continue
                if not node.select_one("div.claim-text, claim-text") and not normalize_text(node.get_text(" ", strip=True)):
                    continue
                seen.add(marker)
                top_level.append(node)
        if top_level:
            return top_level
        fallback = section_node.select("div.claim-text, claim-text, p")
        return fallback or [section_node]

    if section_key == "description":
        root = section_node.body or section_node
        nodes: list[Any] = []
        selected_ids: set[int] = set()
        for node in root.descendants:
            name = getattr(node, "name", None)
            if not name:
                continue
            if any(id(parent) in selected_ids for parent in getattr(node, "parents", []) if getattr(parent, "name", None)):
                continue
            name = name.lower()
            classes = {str(cls).lower() for cls in node.get("class", [])}
            should_select = False
            if name == "heading":
                should_select = True
            elif {"description-paragraph", "description-line"}.intersection(classes):
                should_select = True
            elif name == "p":
                should_select = True
            elif name == "li":
                if node.select_one("div.description-paragraph, div.description-line, p"):
                    should_select = False
                else:
                    should_select = bool(normalize_text(node.get_text(" ", strip=True)))
            if not should_select:
                continue
            nodes.append(node)
            selected_ids.add(id(node))
        return nodes or [section_node]

    selector_map = {
        "abstract": "div.abstract, div.description-paragraph, div.description-line, p",
    }
    nodes = section_node.select(selector_map.get(section_key, "p, div"))
    return nodes or [section_node]


def section_stats_from_blocks(blocks: list[dict[str, Any]]) -> dict[str, int]:
    return {
        "block_count": len(blocks),
        "heading_like_block_count": sum(1 for block in blocks if block.get("is_heading")),
        "inline_image_count": sum(
            1 for block in blocks for token in block.get("tokens", []) if token.get("type") == "inline_image"
        ),
        "missing_object_count": sum(
            1 for block in blocks for token in block.get("tokens", []) if token.get("type") == "missing_object"
        ),
    }


def page_has_drawing_mentions(description_blocks: list[dict[str, Any]]) -> bool:
    for block in description_blocks:
        text = normalize_inline_text(block.get("text", ""))
        if not text:
            continue
        if FIGURE_RE.search(text) or is_drawing_heading_text(text):
            return True
    return False


def build_proxies(proxy_url: str) -> dict[str, str] | None:
    raw = proxy_url.strip()
    if not raw:
        return None
    return {"http": raw, "https": raw}


def normalize_inline_text(raw: str) -> str:
    text = raw.replace("\xa0", " ").replace("\u200b", "")
    for marker in UNICODE_HYPHENS:
        text = text.replace(marker, "-")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[\t\f\v]+", " ", text)
    text = re.sub(r"\s*\n\s*", " ", text)
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"\s+([,.;:!?\)\]\}])", r"\1", text)
    text = re.sub(r"([\(\[\{])\s+", r"\1", text)
    return text.strip()


def normalize_text(raw: str) -> str:
    text = raw.replace("\xa0", " ").replace("\u200b", "")
    for marker in UNICODE_HYPHENS:
        text = text.replace(marker, "-")
    lines: list[str] = []
    for line in text.splitlines():
        cleaned = normalize_inline_text(line)
        if cleaned:
            lines.append(cleaned)

    cleaned_lines: list[str] = []
    noise_exact = {
        "Description",
        "Claims",
        "Abstract",
        "Google Patents",
    }
    for idx, line in enumerate(lines):
        if line in noise_exact and idx == 0:
            continue
        cleaned_lines.append(line)

    return "\n".join(cleaned_lines).strip()


def merge_text_token(tokens: list[dict[str, Any]], raw: str) -> None:
    text = normalize_inline_text(raw)
    if not text:
        return
    if tokens and tokens[-1].get("type") == "text":
        tokens[-1]["text"] = normalize_inline_text(f"{tokens[-1]['text']} {text}")
        return
    tokens.append({"type": "text", "text": text})


def finalize_block_tokens(tokens: list[dict[str, Any]]) -> list[dict[str, Any]]:
    finalized: list[dict[str, Any]] = []
    for token in tokens:
        token_type = token.get("type")
        if token_type == "text":
            merge_text_token(finalized, token.get("text", ""))
            continue
        if token_type == "missing_object" and finalized and finalized[-1].get("type") == "missing_object":
            continue
        finalized.append(token)

    if not finalized:
        return []

    has_text = any(token.get("type") == "text" and token.get("text") for token in finalized)
    for token in finalized:
        if token.get("type") != "missing_object":
            continue
        token["text"] = INLINE_MISSING_PLACEHOLDER if has_text else BLOCK_MISSING_PLACEHOLDER
        token["display_mode"] = "inline" if has_text else "block"
    return finalized


def build_text_from_tokens(tokens: list[dict[str, Any]]) -> str:
    parts = [token.get("text", "") for token in tokens if token.get("type") in {"text", "missing_object"}]
    return normalize_inline_text(" ".join(part for part in parts if part))


def extract_block_tokens(node: Any, base_url: str) -> list[dict[str, Any]]:
    if node is None:
        return []
    if BeautifulSoup is None:
        text = normalize_text(str(node))
        return [{"type": "text", "text": text}] if text else []

    working = BeautifulSoup(str(node), "lxml")
    root = working.body or working
    tokens: list[dict[str, Any]] = []

    def visit(current: Any) -> None:
        if current is None:
            return
        if NavigableString is not None and isinstance(current, NavigableString):
            merge_text_token(tokens, str(current))
            return

        name = getattr(current, "name", None)
        if not name:
            merge_text_token(tokens, str(current))
            return

        name = name.lower()
        classes = set(current.get("class", []))
        if name in {"script", "style"}:
            return
        if "patent-image-not-available" in classes:
            tokens.append({
                "type": "missing_object",
                "source_method": ["span.patent-image-not-available"],
            })
            return
        if name == "div" and "patent-image" in classes:
            link = current.select_one("a[href]")
            image = current.select_one("img[src]")
            url = normalize_url((link.get("href", "") if link else "") or (image.get("src", "") if image else ""), base_url)
            if url:
                source_method = ["section_patent_image"]
                if link is not None:
                    source_method.append("a[href]")
                if image is not None:
                    source_method.append("img[src]")
                tokens.append(
                    {
                        "type": "inline_image",
                        "url": url,
                        "group_key": group_key_from_url(url),
                        "source_method": unique_keep_order(source_method),
                    }
                )
            else:
                tokens.append({
                    "type": "missing_object",
                    "source_method": ["div.patent-image"],
                })
            return
        if name == "br":
            merge_text_token(tokens, " ")
            return

        for child in current.children:
            visit(child)

    for child in root.children:
        visit(child)
    return finalize_block_tokens(tokens)


def section_blocks_to_text(blocks: list[dict[str, Any]], heading: str) -> str:
    lines = [block.get("text", "") for block in blocks if block.get("text")]
    return strip_prefix_heading("\n".join(lines).strip(), heading)


def node_to_clean_text(node: Any, heading: str) -> str:
    if node is None:
        return ""
    tokens = extract_block_tokens(node, "")
    text = build_text_from_tokens(tokens)
    return strip_prefix_heading(text, heading)


def extract_section_blocks(section_node: Any, heading: str, section_key: str, base_url: str) -> list[dict[str, Any]]:
    if section_node is None:
        return []

    nodes = iter_section_block_nodes(section_node, section_key)

    blocks: list[dict[str, Any]] = []
    seen_anchors: set[str] = set()
    for index, node in enumerate(nodes, start=1):
        tokens = extract_block_tokens(node, base_url)
        text = build_text_from_tokens(tokens)
        has_media = any(token.get("type") in {"inline_image", "missing_object"} for token in tokens)
        if not text and not has_media:
            continue
        anchor = node.get("data-codex-anchor") or node.get("id") or f"{section_key}-block-{index:03d}"
        if anchor in seen_anchors:
            anchor = f"{anchor}-{index:03d}"
        node["data-codex-anchor"] = anchor
        seen_anchors.add(anchor)
        blocks.append(
            {
                "section_anchor": anchor,
                "text": text,
                "tokens": tokens,
                "is_heading": is_heading_block(section_key, node, text, tokens),
            }
        )
    return blocks


def strip_prefix_heading(text: str, heading: str) -> str:
    if not text:
        return ""
    lines = text.splitlines()
    if not lines:
        return text
    first = lines[0].strip()
    lower_heading = heading.lower()
    if first.lower() == lower_heading:
        return "\n".join(lines[1:]).strip()
    if first.lower().startswith(lower_heading + " "):
        lines[0] = first[len(heading) :].strip(" :")
        return "\n".join(lines).strip()
    return text


def find_first_node(soup: Any, selectors: list[str]) -> Any | None:
    for selector in selectors:
        node = soup.select_one(selector)
        if node:
            return node
    return None


def find_candidate_nodes(soup: Any, selectors: list[str]) -> list[Any]:
    candidates: list[Any] = []
    seen: set[int] = set()
    for selector in selectors:
        for node in soup.select(selector):
            marker = id(node)
            if marker in seen:
                continue
            seen.add(marker)
            candidates.append(node)
    return candidates


def description_candidate_score(node: Any) -> int:
    if node is None:
        return -1

    text = normalize_text(node.get_text(" ", strip=True))
    lower_text = text.lower()
    parent = node
    while parent is not None and getattr(parent, "name", None):
        itemprop = parent.get("itemprop")
        if itemprop == "events":
            return -1
        classes = {str(cls).lower() for cls in parent.get("class", [])}
        if {"event", "events", "timeline"}.intersection(classes):
            return -1
        parent = parent.parent

    if any(marker in lower_text for marker in DESCRIPTION_METADATA_MARKERS):
        return -1

    score = 0
    if getattr(node, "name", None) == "section":
        score += 100
    if node.select("div.description-paragraph"):
        score += 60
    paragraph_count = len(node.select("p"))
    if paragraph_count >= 3:
        score += 40
    elif paragraph_count:
        score += 10
    for marker in DESCRIPTION_BODY_MARKERS:
        if marker in text:
            score += 20
    score += min(len(text) // 100, 20)
    return score


def looks_invalid_description_text(text: str, reference_text: str = "") -> bool:
    normalized = normalize_text(text)
    if not normalized:
        return True
    lower_text = normalized.lower()
    if any(marker in lower_text for marker in DESCRIPTION_METADATA_MARKERS):
        return True
    has_body_marker = any(marker in normalized for marker in DESCRIPTION_BODY_MARKERS)
    if has_body_marker:
        return False
    if len(normalized) < 120:
        return True
    if reference_text and len(normalized) < max(120, len(reference_text) // 2):
        return True
    return False


def select_section_node(soup: Any, selectors: list[str], section_key: str) -> tuple[Any | None, list[Any]]:
    candidates = find_candidate_nodes(soup, selectors)
    if section_key != "description":
        return (candidates[0] if candidates else None), candidates

    ranked = [
        (description_candidate_score(node), index, node)
        for index, node in enumerate(candidates)
    ]
    ranked = [item for item in ranked if item[0] >= 0]
    ranked.sort(key=lambda item: (-item[0], item[1]))
    ordered = [node for _, _, node in ranked]
    if ordered:
        return ordered[0], ordered
    return (candidates[0] if candidates else None), candidates


def prepare_section_node(node: Any, section_key: str) -> Any:
    if node is None or BeautifulSoup is None:
        return node
    if section_key != "description":
        return node

    fragment = BeautifulSoup(str(node), "lxml")
    for removable in fragment.select("description-of-drawings"):
        removable.decompose()
    return fragment.body or fragment


def extract_section_bundle(
    soup: Any,
    selectors: list[str],
    heading: str,
    section_key: str,
    base_url: str,
    reference_text: str = "",
) -> dict[str, Any]:
    source_node, candidates = select_section_node(soup, selectors, section_key)
    ordered_candidates = candidates if candidates else ([source_node] if source_node is not None else [])

    fallback_bundle = {"text": "", "node": None, "blocks": []}
    for candidate in ordered_candidates:
        working_node = prepare_section_node(candidate, section_key)
        blocks = extract_section_blocks(working_node, heading, section_key, base_url)
        bundle = {
            "text": section_blocks_to_text(blocks, heading),
            "node": working_node,
            "blocks": blocks,
        }
        if fallback_bundle["node"] is None:
            fallback_bundle = bundle
        if section_key != "description":
            return bundle
        if not looks_invalid_description_text(bundle["text"], reference_text):
            return bundle

    return fallback_bundle


def extract_section(soup: Any, selectors: list[str], heading: str) -> str:
    return extract_section_bundle(soup, selectors, heading, "generic", "")["text"]


def normalize_url(url: str, base: str) -> str:
    value = url.strip()
    if not value:
        return ""
    if value.startswith("//"):
        return "https:" + value
    return urljoin(base, value)


def extract_pdf_links(soup: Any, base_url: str, html: str) -> list[str]:
    links: list[str] = []

    def maybe_add(raw: str) -> None:
        href = normalize_url(raw, base_url)
        lower = href.lower()
        if not href or ".pdf" not in lower:
            return
        if PDF_EXT_RE.search(lower) or "patentimages.storage.googleapis.com" in lower:
            links.append(href)

    for node in soup.select("a[href]"):
        maybe_add(node.get("href", ""))

    for node in soup.select("meta[content]"):
        maybe_add(node.get("content", ""))

    for raw in re.findall(r"https://[^\"'\s>]+?\.pdf(?:\?[^\"'\s>]*)?", html, flags=re.IGNORECASE):
        maybe_add(raw)

    return unique_keep_order(links)


def extract_image_candidates(html: str, soup: Any, base_url: str) -> dict[str, set[str]]:
    records: dict[str, set[str]] = {}

    for node in soup.select("a[href]"):
        href = normalize_url(node.get("href", ""), base_url)
        lower = href.lower()
        if "patentimages.storage.googleapis.com" in lower and IMAGE_EXT_RE.search(lower):
            records.setdefault(href, set()).add("a[href]")

    for node in soup.select('img[itemprop="thumbnail"][src]'):
        src = normalize_url(node.get("src", ""), base_url)
        lower = src.lower()
        if "patentimages.storage.googleapis.com" in lower and IMAGE_EXT_RE.search(lower):
            records.setdefault(src, set()).add("img[itemprop=thumbnail]")

    for node in soup.select("img[src]"):
        src = normalize_url(node.get("src", ""), base_url)
        lower = src.lower()
        if "patentimages.storage.googleapis.com" in lower and IMAGE_EXT_RE.search(lower):
            records.setdefault(src, set()).add("img[src]")

    for node in soup.select('meta[itemprop="full"][content]'):
        content = normalize_url(node.get("content", ""), base_url)
        lower = content.lower()
        if "patentimages.storage.googleapis.com" in lower and IMAGE_EXT_RE.search(lower):
            records.setdefault(content, set()).add("meta[itemprop=full]")

    for raw in re.findall(
        r"https://patentimages\.storage\.googleapis\.com/[^\"'\s]+?\.(?:png|jpg|jpeg|webp|tif|tiff)",
        html,
        flags=re.IGNORECASE,
    ):
        records.setdefault(raw, set()).add("script_regex")

    return records


def group_key_from_url(url: str) -> str:
    if not url:
        return ""
    return Path(url.split("?", 1)[0]).stem


def normalize_figure_label(value: str) -> str:
    return unicodedata.normalize("NFKC", value or "").strip()


def split_caption_lines_from_node(node: Any) -> list[str]:
    if node is None:
        return []
    if getattr(node, "name", None) == "figref":
        text = normalize_inline_text(node.get_text(" ", strip=True))
        return [text] if text else []

    figrefs: list[str] = []
    if hasattr(node, "select"):
        for figref in node.select("figref"):
            text = normalize_inline_text(figref.get_text(" ", strip=True))
            if text:
                figrefs.append(text)
    if len(figrefs) > 1:
        return figrefs

    raw_text = ""
    if hasattr(node, "get_text"):
        raw_text = node.get_text("\n", strip=True)
    lines = [normalize_inline_text(part) for part in raw_text.splitlines()]
    return [line for line in lines if line]


def append_drawing_caption_entries(
    captions: list[dict[str, Any]],
    section_anchor: str,
    lines: list[str],
) -> None:
    for line_index, text in enumerate(lines, start=1):
        match = FIGURE_RE.search(text)
        if not match:
            continue
        anchor = section_anchor
        if len(lines) > 1:
            anchor = f"{section_anchor}-{line_index:02d}"
        captions.append(
            {
                "sequence_no": len(captions) + 1,
                "figure_label": normalize_figure_label(match.group(1)),
                "caption_text": text,
                "section_anchor": anchor,
            }
        )


def extract_drawing_captions_from_container(node: Any) -> list[dict[str, Any]]:
    captions: list[dict[str, Any]] = []
    if node is None:
        return captions

    selected_ids: set[int] = set()
    for paragraph in node.descendants:
        name = getattr(paragraph, "name", None)
        if not name:
            continue
        if any(id(parent) in selected_ids for parent in getattr(paragraph, "parents", []) if getattr(parent, "name", None)):
            continue
        name = name.lower()
        classes = {str(cls).lower() for cls in paragraph.get("class", [])}
        should_select = False
        if {"description-paragraph", "description-line"}.intersection(classes):
            should_select = True
        elif name == "p":
            should_select = True
        elif name == "li":
            if paragraph.select_one("div.description-paragraph, div.description-line, p"):
                should_select = False
            else:
                should_select = bool(normalize_text(paragraph.get_text(" ", strip=True)))
        if not should_select:
            continue
        selected_ids.add(id(paragraph))
        section_anchor = paragraph.get("id") or f"drawing-caption-{len(captions) + 1:03d}"
        append_drawing_caption_entries(captions, section_anchor, split_caption_lines_from_node(paragraph))
    return captions


def extract_drawing_captions_from_blocks(blocks: list[dict[str, Any]]) -> list[dict[str, Any]]:
    captions: list[dict[str, Any]] = []
    in_drawing_section = False
    for block in blocks:
        text = normalize_inline_text(block.get("text", ""))
        if not text:
            continue
        if not in_drawing_section:
            if is_drawing_heading_text(text):
                in_drawing_section = True
            continue
        if is_drawing_section_end_text(text):
            break
        if not FIGURE_RE.search(text):
            if captions and block.get("is_heading"):
                break
            continue
        section_anchor = block.get("section_anchor") or f"drawing-caption-{len(captions) + 1:03d}"
        append_drawing_caption_entries(captions, section_anchor, [text])
    return captions


def extract_drawing_captions(soup: Any, description_blocks: list[dict[str, Any]]) -> tuple[list[dict[str, Any]], str]:
    node = soup.select_one("description-of-drawings")
    direct_captions = extract_drawing_captions_from_container(node)
    if direct_captions:
        return direct_captions, "description-of-drawings"

    fallback_captions = extract_drawing_captions_from_blocks(description_blocks)
    if fallback_captions:
        return fallback_captions, "heading-following-list"

    return [], "none"


def extract_gallery_drawing_groups(soup: Any, base_url: str) -> list[dict[str, Any]]:
    groups: list[dict[str, Any]] = []
    for index, node in enumerate(soup.select('li[itemprop="images"]'), start=1):
        thumb_node = node.select_one('img[itemprop="thumbnail"][src]')
        full_node = node.select_one('meta[itemprop="full"][content]')
        thumbnail_url = normalize_url(thumb_node.get("src", ""), base_url) if thumb_node else ""
        full_url = normalize_url(full_node.get("content", ""), base_url) if full_node else ""
        group_key = group_key_from_url(full_url or thumbnail_url)
        if not group_key:
            continue
        groups.append(
            {
                "group_key": group_key,
                "sequence_no": index,
                "source_confidence": "gallery_metadata",
                "full_candidate": {
                    "url": full_url,
                    "source_method": ["meta[itemprop=full]", "gallery_images"],
                }
                if full_url
                else None,
                "thumbnail_candidate": {
                    "url": thumbnail_url,
                    "source_method": ["img[itemprop=thumbnail]", "gallery_images"],
                }
                if thumbnail_url
                else None,
            }
        )
    return groups


def fallback_drawing_groups_from_all_images(
    image_records: dict[str, set[str]], inline_urls: set[str]
) -> list[dict[str, Any]]:
    groups: OrderedDict[str, dict[str, Any]] = OrderedDict()
    for url, methods in image_records.items():
        if url in inline_urls:
            continue
        group_key = group_key_from_url(url)
        if not group_key:
            continue
        group = groups.setdefault(
            group_key,
            {
                "group_key": group_key,
                "sequence_no": len(groups) + 1,
                "source_confidence": "fallback_url_group",
                "full_candidate": None,
                "thumbnail_candidate": None,
            },
        )
        candidate = {"url": url, "source_method": sorted(methods)}
        if "meta[itemprop=full]" in methods and group["full_candidate"] is None:
            group["full_candidate"] = candidate
            continue
        if "img[itemprop=thumbnail]" in methods and group["thumbnail_candidate"] is None:
            group["thumbnail_candidate"] = candidate
            continue
        if group["full_candidate"] is None:
            group["full_candidate"] = candidate
            continue
        if group["thumbnail_candidate"] is None:
            group["thumbnail_candidate"] = candidate
    return list(groups.values())


def find_section_anchor(node: Any, section_key: str, index: int) -> str:
    current = node
    while current is not None and getattr(current, "name", None):
        anchor = current.get("data-codex-anchor") or current.get("id")
        classes = set(current.get("class", []))
        if anchor:
            return anchor
        if classes.intersection({"description-paragraph", "description-line", "claim-text", "abstract", "claim"}):
            anchor = f"{section_key}-block-{index:03d}"
            current["data-codex-anchor"] = anchor
            return anchor
        current = current.parent
    return f"{section_key}-inline-{index:03d}"


def extract_inline_images(section_key: str, blocks: list[dict[str, Any]]) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    images: list[dict[str, Any]] = []
    missing_objects: list[dict[str, Any]] = []
    image_index = 0
    missing_index = 0

    for block in blocks:
        section_anchor = block.get("section_anchor", "")
        for token in block.get("tokens", []):
            token_type = token.get("type")
            if token_type == "inline_image":
                image_index += 1
                images.append(
                    {
                        "section_key": section_key,
                        "section_anchor": section_anchor,
                        "sequence_no": image_index,
                        "url": token.get("url", ""),
                        "group_key": token.get("group_key", ""),
                        "source_method": token.get("source_method", []),
                    }
                )
            elif token_type == "missing_object":
                missing_index += 1
                missing_objects.append(
                    {
                        "section_key": section_key,
                        "section_anchor": section_anchor,
                        "sequence_no": missing_index,
                        "display_text": token.get("text", BLOCK_MISSING_PLACEHOLDER),
                        "source_method": token.get("source_method", ["span.patent-image-not-available"]),
                    }
                )

    return images, missing_objects


def safe_filename(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9._-]+", "_", value).strip("_") or "file"


def fetch_binary(url: str, target: Path, proxies: dict[str, str] | None, timeout_sec: int) -> tuple[bool, int, str]:
    if requests is None:
        return False, 0, "missing_dependency:requests"
    try:
        with requests.get(
            url,
            headers={"User-Agent": USER_AGENT},
            timeout=timeout_sec,
            proxies=proxies,
            stream=True,
        ) as resp:
            if resp.status_code >= 400:
                target.unlink(missing_ok=True)
                return False, 0, f"http_{resp.status_code}"
            target.parent.mkdir(parents=True, exist_ok=True)
            with target.open("wb") as fh:
                for chunk in resp.iter_content(chunk_size=65536):
                    if chunk:
                        fh.write(chunk)
            return True, target.stat().st_size if target.exists() else 0, ""
    except Exception as exc:
        target.unlink(missing_ok=True)
        return False, 0, str(exc)


def render_docx(
    output_path: Path,
    publication_number: str,
    title: str,
    source_url: str,
    sections: dict[str, str],
    section_blocks: dict[str, list[dict[str, Any]]],
    include_flags: dict[str, bool],
    image_manifest: list[dict[str, Any]],
    warnings: list[str],
) -> None:
    try:
        from docx import Document
        from docx.shared import Inches
    except Exception as exc:
        _error(f"missing_dependency: python-docx ({exc})")

    doc = Document()
    configure_docx_styles(doc)
    header = f"{publication_number} - {title}" if title else publication_number
    doc.add_heading(header, level=1)
    doc.add_paragraph(f"Source: {source_url}")
    doc.add_paragraph(f"Generated At: {now_iso()}")

    inline_by_key: dict[tuple[str, str], dict[str, Any]] = {}
    selected_drawings: list[dict[str, Any]] = []
    for item in image_manifest:
        if item.get("image_role") == "inline_formula":
            inline_by_key[(item.get("section_anchor", ""), item.get("url", ""))] = item
        if item.get("selected_for_drawings"):
            selected_drawings.append(item)
    selected_drawings.sort(key=lambda item: item.get("sequence_no", 0))

    def flush_text(parts: list[str]) -> bool:
        cleaned = normalize_inline_text(" ".join(part for part in parts if part))
        parts.clear()
        if not cleaned:
            return False
        doc.add_paragraph(cleaned)
        return True

    block_order = [
        ("abstract", "\u6458\u8981"),
        ("claims", "\u6743\u5229\u8981\u6c42\u4e66"),
        ("description", "\u8bf4\u660e\u4e66"),
    ]

    for key, label in block_order:
        if not include_flags.get(f"include_{key}", True):
            continue
        doc.add_heading(label, level=2)
        blocks = section_blocks.get(key, [])
        rendered_any = False
        if blocks:
            for block in blocks:
                if block.get("is_heading") and block.get("text"):
                    doc.add_heading(block["text"], level=3)
                    rendered_any = True
                    continue
                tokens = block.get("tokens") or ([{"type": "text", "text": block.get("text", "")}] if block.get("text") else [])
                parts: list[str] = []
                block_rendered = False
                for token in tokens:
                    token_type = token.get("type")
                    if token_type in {"text", "missing_object"}:
                        if token.get("text"):
                            parts.append(token["text"])
                        continue
                    if token_type != "inline_image":
                        continue

                    if flush_text(parts):
                        rendered_any = True
                        block_rendered = True

                    entry = inline_by_key.get((block.get("section_anchor", ""), token.get("url", "")))
                    if entry and entry.get("download_status") == "downloaded" and entry.get("local_path"):
                        try:
                            doc.add_picture(entry["local_path"], width=Inches(4.8))
                        except Exception:
                            doc.add_paragraph(INLINE_DOWNLOAD_FAILED_PLACEHOLDER)
                        rendered_any = True
                        block_rendered = True
                        continue

                    doc.add_paragraph(INLINE_DOWNLOAD_FAILED_PLACEHOLDER)
                    rendered_any = True
                    block_rendered = True

                if flush_text(parts):
                    rendered_any = True
                    block_rendered = True

                if not block_rendered and block.get("text"):
                    doc.add_paragraph(block["text"])
                    rendered_any = True
        elif sections.get(key):
            for line in sections[key].splitlines():
                if line.strip():
                    doc.add_paragraph(line.strip())
                    rendered_any = True

        if not rendered_any:
            doc.add_paragraph(CONTENT_MISSING_PLACEHOLDER)

    if include_flags.get("include_drawings", True):
        doc.add_heading("\u9644\u56fe\u8bf4\u660e", level=2)
        if not selected_drawings:
            doc.add_paragraph("\u3010\u672a\u627e\u5230\u53ef\u7528\u9644\u56fe\u3011")
        for item in selected_drawings:
            if item.get("download_status") == "downloaded" and item.get("local_path"):
                try:
                    doc.add_picture(item["local_path"], width=Inches(6.2))
                except Exception:
                    doc.add_paragraph(DRAWING_MISSING_PLACEHOLDER)
            else:
                doc.add_paragraph(DRAWING_MISSING_PLACEHOLDER)

            title_line = f"\u56fe{item.get('sequence_no', '')}"
            caption_text = normalize_inline_text(item.get("caption_text", ""))
            if caption_text:
                if caption_text.startswith(title_line):
                    title_line = caption_text
                else:
                    title_line = f"{title_line} {caption_text}"
            if item.get("fallback_reason"):
                title_line = f"{title_line}\uff08\u7f29\u7565\u56fe\u56de\u9000\uff09"
            if item.get("download_status") != "downloaded":
                title_line = f"{title_line}\uff08\u56fe\u7247\u7f3a\u5931\uff09"
            doc.add_paragraph(title_line)

    if warnings:
        doc.add_heading("Warnings", level=2)
        for warning in warnings:
            doc.add_paragraph(warning, style="List Bullet")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def process_publication(
    publication_number: str,
    output_root: Path,
    file_formats: str,
    proxy_url: str,
    timeout_sec: int,
    include_flags: dict[str, bool],
) -> dict[str, Any]:
    if requests is None:
        _error("missing_dependency:requests")
    if BeautifulSoup is None:
        _error("missing_dependency:bs4")

    source_url = PATENT_BASE_URL.format(publication_number=publication_number)
    proxies = build_proxies(proxy_url)
    warnings: list[str] = []

    try:
        response = requests.get(
            source_url,
            headers={"User-Agent": USER_AGENT},
            timeout=timeout_sec,
            proxies=proxies,
            allow_redirects=True,
        )
    except Exception as exc:
        return {
            "publication_number": publication_number,
            "ok": False,
            "source_url": source_url,
            "warnings": [f"fetch_failed:{exc}"],
        }

    if response.status_code >= 400:
        return {
            "publication_number": publication_number,
            "ok": False,
            "source_url": source_url,
            "final_url": response.url,
            "warnings": [f"fetch_http_{response.status_code}"],
        }

    response.encoding = response.apparent_encoding or response.encoding or "utf-8"
    html = response.text
    soup = BeautifulSoup(html, "lxml")

    patent_dir = output_root / publication_number
    images_dir = patent_dir / "images"
    patent_dir.mkdir(parents=True, exist_ok=True)

    raw_html_path = patent_dir / "raw_html.html"
    raw_html_path.write_text(html, encoding="utf-8")

    title = ""
    meta_title = soup.select_one('meta[name="DC.title"]')
    if meta_title and meta_title.get("content"):
        title = str(meta_title.get("content")).strip()
    if not title:
        title_node = soup.select_one("title")
        if title_node:
            title = normalize_text(title_node.get_text(" ", strip=True))
            title = title.replace("- Google Patents", "").strip()

    section_specs = {
        "abstract": (
            ['[itemprop="abstract"]', "section.abstract", "div.abstract"],
            "Abstract",
        ),
        "claims": (
            ['[itemprop="claims"]', "section#claims", "section.claims"],
            "Claims",
        ),
        "description": (
            ['section[itemprop="description"]', "section#description", "section.description", '[itemprop="description"]'],
            "Description",
        ),
    }
    section_bundles: dict[str, dict[str, Any]] = {}
    for key, (selectors, heading) in section_specs.items():
        reference_text = section_bundles.get("abstract", {}).get("text", "") if key == "description" else ""
        section_bundles[key] = extract_section_bundle(
            soup,
            selectors,
            heading,
            key,
            response.url,
            reference_text=reference_text,
        )
    sections = {key: bundle["text"] for key, bundle in section_bundles.items()}
    section_blocks = {key: bundle["blocks"] for key, bundle in section_bundles.items()}

    pdf_links = extract_pdf_links(soup, response.url, html)
    image_records = extract_image_candidates(html, soup, response.url)
    drawing_captions, drawing_caption_source = extract_drawing_captions(soup, section_bundles["description"]["blocks"])
    source_office = detect_source_office(publication_number)

    inline_candidates: list[dict[str, Any]] = []
    inline_missing_candidates: list[dict[str, Any]] = []
    for key in ("abstract", "claims", "description"):
        if not include_flags.get(f"include_{key}", True):
            continue
        section_inline, section_missing = extract_inline_images(key, section_bundles[key]["blocks"])
        inline_candidates.extend(section_inline)
        inline_missing_candidates.extend(section_missing)
    inline_urls = {item["url"] for item in inline_candidates if item.get("url")}

    drawing_groups = extract_gallery_drawing_groups(soup, response.url)
    if not drawing_groups:
        drawing_groups = fallback_drawing_groups_from_all_images(image_records, inline_urls)
        if drawing_groups:
            warnings.append("drawing_groups_from_fallback_sources")

    drawing_targets = drawing_captions
    if not sections.get("abstract"):
        warnings.append("abstract_not_exposed_on_page")

    if not drawing_targets and drawing_groups:
        warnings.append("drawing_captions_not_found")
        drawing_targets = [
            {
                "sequence_no": index,
                "figure_label": str(index),
                "caption_text": "",
                "section_anchor": f"drawing-caption-{index:03d}",
            }
            for index in range(1, len(drawing_groups) + 1)
        ]
    elif drawing_targets and not drawing_groups:
        warnings.append("google_drawings_not_exposed_on_page")
    elif not drawing_targets and not drawing_groups and page_has_drawing_mentions(section_bundles["description"]["blocks"]):
        warnings.append("drawing_captions_not_found")

    if drawing_captions and drawing_groups and len(drawing_captions) != len(drawing_groups):
        warnings.append(
            f"drawing_caption_image_count_mismatch:{len(drawing_captions)}:{len(drawing_groups)}"
        )

    image_manifest: list[dict[str, Any]] = []
    manifest_counter = {"value": 0}
    download_counter = {"value": 0}
    download_cache: dict[str, dict[str, Any]] = {}
    handled_urls: set[str] = set()
    include_images = include_flags.get("include_drawings", True)

    def ensure_download(url: str, should_download: bool, skipped_status: str) -> dict[str, Any]:
        if not url:
            return {
                "download_status": "missing_url",
                "local_path": "",
                "file_size": 0,
                "error": "missing_url",
            }
        if not should_download:
            return {
                "download_status": skipped_status,
                "local_path": "",
                "file_size": 0,
                "error": "",
            }
        if url in download_cache:
            return download_cache[url]
        download_counter["value"] += 1
        suffix = Path(url.split("?", 1)[0]).suffix or ".png"
        name = f"{download_counter['value']:03d}_{safe_filename(Path(url.split('?', 1)[0]).name)}"
        if not name.lower().endswith(suffix.lower()):
            name = f"{name}{suffix}"
        local_path = images_dir / name
        ok, size, error = fetch_binary(url, local_path, proxies, timeout_sec)
        download_cache[url] = {
            "download_status": "downloaded" if ok else "failed",
            "local_path": str(local_path) if ok else "",
            "file_size": size,
            "error": error,
        }
        return download_cache[url]

    def append_manifest_entry(**kwargs: Any) -> dict[str, Any]:
        manifest_counter["value"] += 1
        entry = {
            "index": manifest_counter["value"],
            "url": kwargs.get("url", ""),
            "source_method": kwargs.get("source_method", []),
            "download_status": kwargs.get("download_status", "recorded_only"),
            "local_path": kwargs.get("local_path", ""),
            "file_size": kwargs.get("file_size", 0),
            "error": kwargs.get("error", ""),
            "image_role": kwargs.get("image_role", "other"),
            "group_key": kwargs.get("group_key", ""),
            "selected_for_drawings": kwargs.get("selected_for_drawings", False),
            "fallback_reason": kwargs.get("fallback_reason", ""),
            "section_anchor": kwargs.get("section_anchor", ""),
            "sequence_no": kwargs.get("sequence_no", 0),
            "caption_text": kwargs.get("caption_text", ""),
            "figure_label": kwargs.get("figure_label", ""),
            "section_key": kwargs.get("section_key", ""),
            "source_confidence": kwargs.get("source_confidence", ""),
        }
        image_manifest.append(entry)
        return entry

    selected_drawings: list[dict[str, Any]] = []
    pair_count = min(len(drawing_targets), len(drawing_groups)) if drawing_targets and drawing_groups else 0
    for index in range(pair_count):
        caption = drawing_targets[index]
        group = drawing_groups[index]
        full_candidate = group.get("full_candidate") or {}
        thumbnail_candidate = group.get("thumbnail_candidate") or {}
        full_url = full_candidate.get("url", "")
        thumbnail_url = thumbnail_candidate.get("url", "")
        if full_url:
            handled_urls.add(full_url)
        if thumbnail_url:
            handled_urls.add(thumbnail_url)

        full_info = ensure_download(full_url, include_images, "skipped") if full_url else {
            "download_status": "missing_candidate",
            "local_path": "",
            "file_size": 0,
            "error": "missing_full_candidate",
        }
        thumbnail_info = {
            "download_status": "not_attempted",
            "local_path": "",
            "file_size": 0,
            "error": "",
        }
        fallback_reason = ""
        selected_url = ""

        if full_info["download_status"] == "downloaded":
            selected_url = full_url
        else:
            thumbnail_info = ensure_download(thumbnail_url, include_images, "skipped") if thumbnail_url else {
                "download_status": "missing_candidate",
                "local_path": "",
                "file_size": 0,
                "error": "missing_thumbnail_candidate",
            }
            if thumbnail_info["download_status"] == "downloaded":
                selected_url = thumbnail_url
                fallback_reason = f"hd_unavailable:{full_info.get('error') or full_info.get('download_status', '')}"

        full_entry = None
        if full_url:
            full_entry = append_manifest_entry(
                url=full_url,
                source_method=full_candidate.get("source_method", []),
                download_status=full_info.get("download_status", "missing_candidate"),
                local_path=full_info.get("local_path", ""),
                file_size=full_info.get("file_size", 0),
                error=full_info.get("error", ""),
                image_role="drawing_hd",
                group_key=group.get("group_key", ""),
                selected_for_drawings=selected_url == full_url,
                section_anchor=caption.get("section_anchor", ""),
                sequence_no=caption.get("sequence_no", index + 1),
                caption_text=caption.get("caption_text", ""),
                figure_label=caption.get("figure_label", ""),
                source_confidence=group.get("source_confidence", ""),
            )

        thumbnail_entry = None
        if thumbnail_url:
            thumbnail_entry = append_manifest_entry(
                url=thumbnail_url,
                source_method=thumbnail_candidate.get("source_method", []),
                download_status=thumbnail_info.get("download_status", "missing_candidate"),
                local_path=thumbnail_info.get("local_path", ""),
                file_size=thumbnail_info.get("file_size", 0),
                error=thumbnail_info.get("error", ""),
                image_role="drawing_thumbnail",
                group_key=group.get("group_key", ""),
                selected_for_drawings=selected_url == thumbnail_url,
                fallback_reason=fallback_reason if selected_url == thumbnail_url else "",
                section_anchor=caption.get("section_anchor", ""),
                sequence_no=caption.get("sequence_no", index + 1),
                caption_text=caption.get("caption_text", ""),
                figure_label=caption.get("figure_label", ""),
                source_confidence=group.get("source_confidence", ""),
            )

        chosen_entry = full_entry if selected_url == full_url else thumbnail_entry
        if chosen_entry is not None:
            selected_drawings.append(chosen_entry)
        elif include_images:
            warnings.append(f"drawing_download_failed:{group.get('group_key', '')}")

    if drawing_groups and pair_count < len(drawing_groups):
        for leftover in drawing_groups[pair_count:]:
            for candidate_name in ("full_candidate", "thumbnail_candidate"):
                candidate = leftover.get(candidate_name) or {}
                url = candidate.get("url", "")
                if not url:
                    continue
                handled_urls.add(url)
                append_manifest_entry(
                    url=url,
                    source_method=candidate.get("source_method", []),
                    download_status="recorded_only",
                    image_role="other",
                    group_key=leftover.get("group_key", ""),
                    source_confidence=leftover.get("source_confidence", ""),
                )

    for candidate in inline_candidates:
        handled_urls.add(candidate["url"])
        info = ensure_download(candidate["url"], include_images, "skipped")
        append_manifest_entry(
            url=candidate["url"],
            source_method=candidate["source_method"],
            download_status=info.get("download_status", "skipped"),
            local_path=info.get("local_path", ""),
            file_size=info.get("file_size", 0),
            error=info.get("error", ""),
            image_role="inline_formula",
            group_key=candidate.get("group_key", ""),
            section_anchor=candidate.get("section_anchor", ""),
            sequence_no=candidate.get("sequence_no", 0),
            section_key=candidate.get("section_key", ""),
            source_confidence="section_dom",
        )

    for candidate in inline_missing_candidates:
        append_manifest_entry(
            url="",
            source_method=candidate.get("source_method", ["span.patent-image-not-available"]),
            download_status="not_available_on_page",
            local_path="",
            file_size=0,
            error="",
            image_role="inline_formula_missing",
            group_key="",
            section_anchor=candidate.get("section_anchor", ""),
            sequence_no=candidate.get("sequence_no", 0),
            section_key=candidate.get("section_key", ""),
            caption_text=candidate.get("display_text", ""),
            source_confidence="section_dom",
        )

    for url, source_set in image_records.items():
        if url in handled_urls:
            continue
        append_manifest_entry(
            url=url,
            source_method=sorted(source_set),
            download_status="recorded_only",
            image_role="other",
            group_key=group_key_from_url(url),
            source_confidence="residual_trace",
        )

    if include_images and not selected_drawings:
        warnings.append("no_drawings_found_after_full_and_thumbnail_attempt")

    image_manifest_path = patent_dir / "image_manifest.json"
    image_manifest_path.write_text(json.dumps(image_manifest, ensure_ascii=False, indent=2), encoding="utf-8")

    pdf_source = "none"
    pdf_path = ""
    # 改进：不管 file_formats 是什么，只要有 PDF 链接就尝试下载
    # 除非用户明确指定只想要 docx（可以通过单独的标志位控制，这里简化处理）
    if file_formats in {"pdf", "both"} and pdf_links:
        # 始终尝试下载 PDF（如果存在）
        target_pdf = patent_dir / f"{publication_number}.pdf"
        pdf_errors: list[str] = []
        for pdf_link in pdf_links:
            ok, _, err = fetch_binary(pdf_link, target_pdf, proxies, timeout_sec)
            if ok:
                pdf_source = "google"
                pdf_path = str(target_pdf)
                break
            pdf_errors.append(f"{pdf_link}|{err}")
        if not pdf_path:
            reason = pdf_errors[0].split("|", 1)[1] if pdf_errors else "unknown"
            warnings.append(f"google_pdf_download_failed:{reason}")
    else:
        warnings.append("google_pdf_not_found")

    warnings = unique_keep_order(warnings)

    docx_path = ""
    if file_formats in {"docx", "both"}:
        docx_target = patent_dir / f"{publication_number}.docx"
        render_docx(
            docx_target,
            publication_number=publication_number,
            title=title,
            source_url=source_url,
            sections=sections,
            section_blocks=section_blocks,
            include_flags=include_flags,
            image_manifest=image_manifest,
            warnings=warnings,
        )
        docx_path = str(docx_target)

    parsed_payload = {
        "publication_number": publication_number,
        "source_office": source_office,
        "title": title,
        "source_url": source_url,
        "final_url": response.url,
        "status_code": response.status_code,
        "sections": sections,
        "section_stats": {key: section_stats_from_blocks(blocks) for key, blocks in section_blocks.items()},
        "pdf_links": pdf_links,
        "image_total": len(image_manifest),
        "drawing_captions": drawing_captions,
        "drawing_caption_total": len(drawing_captions),
        "drawing_caption_source": drawing_caption_source,
        "drawing_group_total": len(drawing_groups),
        "inline_object_total": len(inline_candidates),
        "inline_missing_total": len(inline_missing_candidates),
    }
    parsed_json_path = patent_dir / "parsed.json"
    parsed_json_path.write_text(json.dumps(parsed_payload, ensure_ascii=False, indent=2), encoding="utf-8")

    metadata = {
        "publication_number": publication_number,
        "source_office": source_office,
        "timestamp": now_iso(),
        "source_url": source_url,
        "final_url": response.url,
        "used_proxy": proxy_url,
        "file_formats": file_formats,
        "pdf_source": pdf_source,
        "image_count": len(image_manifest),
        "drawing_total": len(drawing_groups),
        "drawing_caption_total": len(drawing_captions),
        "drawing_caption_source": drawing_caption_source,
        "drawing_selected": len(selected_drawings),
        "drawing_hd_selected": sum(1 for item in selected_drawings if item.get("image_role") == "drawing_hd"),
        "drawing_thumbnail_fallback": sum(
            1 for item in selected_drawings if item.get("image_role") == "drawing_thumbnail"
        ),
        "inline_object_total": len(inline_candidates),
        "inline_missing_total": len(inline_missing_candidates),
        "warnings": warnings,
        "contract": {
            "en_suffix_disabled": True,
            "online_required": True,
        },
    }
    run_metadata_path = patent_dir / "run_metadata.json"
    run_metadata_path.write_text(json.dumps(metadata, ensure_ascii=False, indent=2), encoding="utf-8")

    return {
        "publication_number": publication_number,
        "ok": True,
        "source_url": source_url,
        "final_url": response.url,
        "raw_html_path": str(raw_html_path),
        "parsed_json_path": str(parsed_json_path),
        "image_manifest_path": str(image_manifest_path),
        "run_metadata_path": str(run_metadata_path),
        "docx_path": docx_path,
        "pdf_path": pdf_path,
        "warnings": warnings,
    }


def run_patent_mode(args: argparse.Namespace) -> dict[str, Any]:
    payload = load_payload(args)
    context = flatten_context(payload)

    direct_publications = as_list(args.publication_numbers)
    context_publications = as_list(context.get("publication_numbers"))
    inline_publication = as_list(context.get("publication_number"))
    publication_numbers = unique_keep_order(direct_publications + context_publications + inline_publication)

    if not publication_numbers:
        _error("publication_numbers_required")

    file_formats = as_str(context.get("file_formats") or "both").lower()
    if file_formats not in {"docx", "pdf", "both"}:
        file_formats = "both"

    proxy_url = as_str(args.proxy_url or context.get("proxy_url") or os.environ.get("HTTPS_PROXY") or os.environ.get("HTTP_PROXY") or "")
    timeout_sec = as_int(args.timeout_sec or context.get("request_timeout_sec"), 30)
    timeout_sec = max(5, min(timeout_sec, 120))

    include_flags = {
        "include_abstract": as_bool(context.get("include_abstract"), True),
        "include_claims": as_bool(context.get("include_claims"), True),
        "include_description": as_bool(context.get("include_description"), True),
        "include_drawings": as_bool(context.get("include_drawings"), True),
    }

    output_root = Path(args.output_dir).expanduser().resolve() if args.output_dir else Path.cwd() / "patent_outputs"
    output_root.mkdir(parents=True, exist_ok=True)

    results: list[dict[str, Any]] = []
    for publication_number in publication_numbers:
        results.append(
            process_publication(
                publication_number=publication_number,
                output_root=output_root,
                file_formats=file_formats,
                proxy_url=proxy_url,
                timeout_sec=timeout_sec,
                include_flags=include_flags,
            )
        )

    text_table_path, text_table_row_count = export_patent_text_table(results, output_root)

    return {
        "ok": True,
        "mode": "detail",
        "output_dir": str(output_root),
        "processed_count": len(results),
        "publication_numbers": publication_numbers,
        "text_table_path": text_table_path,
        "text_table_row_count": text_table_row_count,
        "results": results,
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="Export Google Patents detail artifacts")
    parser.add_argument("--mode", default="patent", choices=["patent"])
    parser.add_argument("--output-dir", default=None, help="Output directory for detail artifacts.")
    parser.add_argument("--answers", help="JSON payload from intake/file_ingest outputs")
    parser.add_argument("--answers-file", help="Path to JSON payload from intake/file_ingest outputs")
    parser.add_argument("--context", help="Inline JSON context payload")
    parser.add_argument("--context-file", help="Path to JSON context payload")
    parser.add_argument("--publication-numbers", help="Comma/newline separated publication numbers")
    parser.add_argument("--proxy-url", help="HTTP/HTTPS proxy URL")
    parser.add_argument("--timeout-sec", type=int, help="Request timeout seconds")

    args = parser.parse_args()
    result = run_patent_mode(args)
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()




