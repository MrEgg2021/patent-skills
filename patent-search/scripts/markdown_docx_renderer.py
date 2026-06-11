#!/usr/bin/env python3
from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Callable


def _strip_inline_markdown(text: str) -> str:
    cleaned = re.sub(r'!\[[^\]]*\]\([^\)]*\)', '[image]', text)
    cleaned = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', cleaned)
    cleaned = re.sub(r'`([^`]+)`', r'\1', cleaned)
    cleaned = re.sub(r'\*\*([^*]+)\*\*', r'\1', cleaned)
    cleaned = re.sub(r'__([^_]+)__', r'\1', cleaned)
    cleaned = re.sub(r'\*([^*]+)\*', r'\1', cleaned)
    cleaned = re.sub(r'_([^_]+)_', r'\1', cleaned)
    cleaned = re.sub(r'~~([^~]+)~~', r'\1', cleaned)
    return cleaned


def render_markdown_to_docx(
    markdown_text: str,
    output_path: str | Path,
    configure_document: Callable[[Any], None] | None = None,
) -> None:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError(f'missing_dependency: python-docx ({exc})') from exc

    document = Document()
    if configure_document is not None:
        configure_document(document)

    in_code_block = False
    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip('\n')
        stripped = line.strip()
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            paragraph = document.add_paragraph(style='Normal')
            paragraph.add_run(line)
            continue

        if stripped == '':
            document.add_paragraph('')
            continue

        heading_match = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 3)
            document.add_paragraph(_strip_inline_markdown(heading_match.group(2)), style=f'Heading {level}')
            continue

        bullet_match = re.match(r'^[-*]\s+(.*)$', stripped)
        if bullet_match:
            document.add_paragraph(_strip_inline_markdown(bullet_match.group(1)), style='List Bullet')
            continue

        number_match = re.match(r'^\d+[\.)]\s+(.*)$', stripped)
        if number_match:
            document.add_paragraph(_strip_inline_markdown(number_match.group(1)), style='List Number')
            continue

        document.add_paragraph(_strip_inline_markdown(line), style='Normal')

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))