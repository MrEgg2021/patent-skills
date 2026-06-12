#!/usr/bin/env python3
"""技术交底书审核报告 Word 生成器。

读取结构化 JSON，输出三段式审核报告 .docx：
  1. 领域识别与整体评估
  2. 逐条审查记录（表格）
  3. 发明人补充确认清单（Part A 宏观 / Part B 微观）

用法：
  python render_review_docx.py --data report.json --output 报告.docx

JSON 结构见同目录 references/review_report_schema.md 或本文件末尾示例。
缺失字段一律优雅跳过，不报错。
"""
from __future__ import annotations

import argparse
import json
import os
import sys

# 跨平台：确保非 ASCII 输出在 Windows GBK 控制台不崩溃
try:
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass

try:
    import docx  # noqa: F401
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print(
        json.dumps(
            {"error": "python-docx 未安装，请先运行: pip install python-docx"},
            ensure_ascii=False,
        )
    )
    raise SystemExit(2)


FONT = "Microsoft YaHei"
HEADER_FILL = "2B579A"  # 表头深蓝底
BORDER_COLOR = "999999"

# 级别 → (显示文字, 字体颜色)。兼容 emoji 与纯文字两种输入写法。
LEVEL_MAP = {
    "🔴": ("🔴 必须补充", RGBColor(0xC0, 0x00, 0x00)),
    "必须补充": ("🔴 必须补充", RGBColor(0xC0, 0x00, 0x00)),
    "🔵": ("🔵 建议补充", RGBColor(0x1F, 0x4E, 0x99)),
    "建议补充": ("🔵 建议补充", RGBColor(0x1F, 0x4E, 0x99)),
    "✅": ("✅ 通过", RGBColor(0x2E, 0x7D, 0x32)),
    "通过": ("✅ 通过", RGBColor(0x2E, 0x7D, 0x32)),
}


def _level(raw):
    """归一化级别字段，返回 (显示文字, 颜色)。未知值原样显示、黑字。"""
    if not raw:
        return ("", RGBColor(0x00, 0x00, 0x00))
    key = str(raw).strip()
    for k, v in LEVEL_MAP.items():
        if key.startswith(k):
            return v
    return (key, RGBColor(0x00, 0x00, 0x00))


def set_font(run, size=10.5, bold=False, color=None):
    """设置 run 字体：西文+东亚双设，避免中文不生效。"""
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), FONT)


def set_cell_border(cell, color=BORDER_COLOR, sz=4):
    """四边加边框。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge in ("top", "left", "bottom", "right"):
        el = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="single" '
            f'w:sz="{sz}" w:space="0" w:color="{color}"/>'
        )
        tcBorders.append(el)
    tcPr.append(tcBorders)


def shade_cell(cell, fill=HEADER_FILL):
    """单元格底色。"""
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shd)


def set_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)


def add_heading(doc, text, size=14, bold=True, color=None):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p


def add_para(doc, label, value, label_bold=True):
    """加一段「标签：内容」。label 可为空时只写 value。"""
    p = doc.add_paragraph()
    if label:
        r1 = p.add_run(label)
        set_font(r1, size=10.5, bold=label_bold)
    if value:
        r2 = p.add_run(value)
        set_font(r2, size=10.5, bold=False)
    return p


def _set_cell_text(cell, text, bold=False, color=None, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text if text else "")
    set_font(run, size=10, bold=bold, color=color)
    set_cell_border(cell)
    if fill:
        shade_cell(cell, fill)


def render_assessment(doc, data):
    """第 1 段：领域识别与整体评估。"""
    add_heading(doc, "一、领域识别与整体评估", size=14)
    domain = data.get("domain", "")
    if domain:
        add_para(doc, "识别领域：", domain)
    overall = data.get("overall_assessment", "")
    if overall:
        add_para(doc, "整体评估：", overall)
    doc.add_paragraph()


def render_review_records(doc, data):
    """第 2 段：逐条审查记录（表格）。"""
    records = data.get("review_records", []) or []
    if not records:
        return
    add_heading(doc, "二、逐条审查记录", size=14)
    table = doc.add_table(rows=1, cols=4)
    table.autofit = True
    headers = ["规则", "级别", "现状", "问题与影响"]
    for i, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], h, bold=True,
                       color=RGBColor(0xFF, 0xFF, 0xFF), fill=HEADER_FILL)
    for rec in records:
        row = table.add_row().cells
        rule = rec.get("rule", "")
        name = rec.get("name", "")
        rule_text = f"{rule} {name}".strip()
        disp, col = _level(rec.get("level", ""))
        status = rec.get("status", "")
        problem = rec.get("problem", "")
        impact = rec.get("impact", "")
        pi = problem
        if impact:
            pi = f"{problem}\n影响：{impact}" if problem else f"影响：{impact}"
        _set_cell_text(row[0], rule_text, bold=True)
        _set_cell_text(row[1], disp, color=col)
        _set_cell_text(row[2], status)
        _set_cell_text(row[3], pi)
    doc.add_paragraph()


def _render_confirm_item(doc, item, prefix):
    """渲染一个补充确认项（Part A / Part B 共用）。"""
    title = item.get("title", "")
    add_heading(doc, f"{prefix} {title}".strip(), size=11.5, bold=True)
    origin = item.get("origin_ref", "")
    if origin:
        p = add_para(doc, "原文定位：", "")
        q = doc.add_paragraph()
        q.paragraph_format.left_indent = Cm(0.74)
        run = q.add_run(origin)
        set_font(run, size=10, bold=False, color=RGBColor(0x55, 0x55, 0x55))
    desc = item.get("description", "")
    if desc:
        add_para(doc, "问题描述：", desc)
    level = item.get("level", "")
    if level:
        disp, col = _level(level)
        p = doc.add_paragraph()
        r1 = p.add_run("严重级别：")
        set_font(r1, size=10.5, bold=True)
        r2 = p.add_run(disp)
        set_font(r2, size=10.5, bold=False, color=col)
    impact = item.get("impact", "")
    if impact:
        add_para(doc, "为什么影响撰写：", impact)
    options = item.get("options", []) or []
    if options:
        add_para(doc, "请发明人确认/补充：", "")
        for opt in options:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(opt)
            set_font(run, size=10.5, bold=False)
    doc.add_paragraph()


def render_confirm_list(doc, data):
    """第 3 段：发明人补充确认清单（Part A / Part B）。"""
    part_a = data.get("part_a", []) or []
    part_b = data.get("part_b", []) or []
    if not part_a and not part_b:
        return
    add_heading(doc, "三、发明人补充确认清单", size=14)
    if part_a:
        add_heading(doc, "Part A：整体性 / 宏观问题", size=12,
                    color=RGBColor(0x2B, 0x57, 0x9A))
        for i, item in enumerate(part_a, 1):
            _render_confirm_item(doc, item, f"A-{i}.")
    if part_b:
        add_heading(doc, "Part B：具体点 / 微观问题", size=12,
                    color=RGBColor(0x2B, 0x57, 0x9A))
        for i, item in enumerate(part_b, 1):
            _render_confirm_item(doc, item, f"B-{i}.")


def build_report(data, output_path):
    doc = Document()
    set_margins(doc)
    # 标题
    title_text = data.get("invention_title", "技术交底书")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"《{title_text}》技术交底书审核报告")
    set_font(run, size=18, bold=True)
    doc.add_paragraph()

    render_assessment(doc, data)
    render_review_records(doc, data)
    render_confirm_list(doc, data)

    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description="技术交底书审核报告 Word 生成器")
    parser.add_argument("--data", required=True, help="审核结果 JSON 字符串或文件路径")
    parser.add_argument("--output", required=True, help="输出 .docx 路径")
    args = parser.parse_args()

    raw = args.data
    if os.path.exists(raw):
        with open(raw, encoding="utf-8") as f:
            data = json.load(f)
    else:
        data = json.loads(raw)

    out = build_report(data, args.output)
    print(json.dumps({"ok": True, "output": os.path.abspath(out)}, ensure_ascii=False))


if __name__ == "__main__":
    main()
