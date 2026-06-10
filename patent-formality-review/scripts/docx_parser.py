#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCX分节解析器 — 专利申请初稿形式审查

通过分节符+页眉识别专利DOCX的四部分：
  - 说明书摘要
  - 权利要求书
  - 说明书
  - 说明书附图

提取各部分文本、附图标记、图片、表格、公式等结构化数据。
"""

# 跨平台：确保非 ASCII（中文/emoji）输出在 Windows GBK 控制台不崩溃
import sys as _sys
try:
    _sys.stdout.reconfigure(encoding='utf-8')
    _sys.stderr.reconfigure(encoding='utf-8')
except Exception:
    pass
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("错误：需要 python-docx 库。请运行：pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ──────────────────────────────────────────────
# 数据结构
# ──────────────────────────────────────────────

@dataclass
class ImageInfo:
    """嵌入图片信息"""
    index: int               # 在段落中的顺序
    data: bytes              # 图片二进制数据
    width_pt: float          # 宽度(磅)
    height_pt: float         # 高度(磅)
    format: str              # 格式(png/jpeg/emf/wmf等)
    is_color: Optional[bool] = None   # 是否彩色(None=未检测)


@dataclass
class SectionData:
    """一个分节的数据"""
    section_type: str           # '摘要'|'权利要求书'|'说明书'|'说明书附图'|'未知'
    header_text: str            # 页眉文本
    paragraphs: list = field(default_factory=list)       # 段落文本列表
    para_index_range: tuple = (0, 0)  # (起始段落全局序号, 结束段落全局序号)
    figure_marks: list = field(default_factory=list)     # 附图标记列表(字符串)
    images: list = field(default_factory=list)           # ImageInfo列表
    tables: list = field(default_factory=list)           # 表格内容[list[list[str]]]
    has_omath: bool = False     # 是否含OOXML公式
    raw_section: object = None  # python-docx section对象(内部用，可选)

    @property
    def full_text(self) -> str:
        """获取完整文本"""
        return '\n'.join(self.paragraphs)


@dataclass
class ParsedDocument:
    """解析后的完整文档"""
    file_path: str = ""
    patent_type: str = '未知'     # '发明'|'实用新型'|'未知'
    invention_name: str = ''      # 发明名称
    sections: dict = field(default_factory=dict)  # section_type -> SectionData
    warnings: list = field(default_factory=list)  # 解析警告
    all_paragraphs: list = field(default_factory=list)  # 全部段落文本(含索引)


# ──────────────────────────────────────────────
# 解析器主类
# ──────────────────────────────────────────────

class DocxSectionParser:
    """通过分节符+页眉识别专利DOCX的四部分"""

    # 页眉关键词匹配规则（兼容空格/全角/半角）
    HEADER_PATTERNS = {
        '摘要':     re.compile(r'摘\s*要'),
        '权利要求书': re.compile(r'权\s*利\s*要\s*求'),
        '说明书':   re.compile(r'说\s*明\s*书'),
        '说明书附图': re.compile(r'附\s*图'),
    }

    # 正文标题匹配规则
    TITLE_PATTERNS = {
        '摘要':     re.compile(r'^说?\s*明?\s*书?\s*摘\s*要$'),
        '权利要求书': re.compile(r'^权\s*利\s*要\s*求\s*书$'),
        '说明书':   re.compile(r'^说\s*明\s*书$'),
        '说明书附图': re.compile(r'^说?\s*明?\s*书?\s*附\s*图$'),
    }

    # 五部分标题（用于说明书内部分析）
    SPEC_SUBHEADINGS = [
        re.compile(r'^技\s*术\s*领\s*域$'),
        re.compile(r'^背\s*景\s*技\s*术$'),
        re.compile(r'^(发\s*明|实\s*用\s*新\s*型)\s*内\s*容$'),
        re.compile(r'^附\s*图\s*说\s*明$'),
        re.compile(r'^具\s*体\s*实\s*施\s*方\s*式$'),
    ]

    # 专利类型用语
    PATENT_TYPE_PATTERNS = {
        '实用新型': re.compile(r'本\s*实\s*用\s*新\s*型'),
        '发明':     re.compile(r'本\s*发\s*明'),
    }

    # 附图标记匹配（括号内数字）
    FIGURE_MARK_IN_BRACKETS = re.compile(r'（(\d{1,4})）|\((\d{1,4})\)')
    # 正文中的附图标记（紧跟在中文字符后的2-4位数字，排除常见度量单位与量词）
    # 排除词包括：年月日号个倍%％米秒克瓦伏安兆至人毫秒阈值数值上限下限比率比值比例程度次被削从还每求元维关
    FIGURE_MARK_IN_TEXT = re.compile(
        r'[\u4e00-\u9fff](\d{2,4})'
        r'(?!['
        r'\d年月日号个倍%％米秒克瓦伏安兆'
        r'至人毫秒阈值数值上下限率比程度幅围内时期间隔帧页张项条段句字节位点级次被削从还每求元维关]'  # 追加排除量词/描述词/动词
        r')'
    )

    def __init__(self, docx_path: str):
        self.docx_path = Path(docx_path)
        if not self.docx_path.exists():
            raise FileNotFoundError(f"文件不存在: {docx_path}")
        if self.docx_path.suffix.lower() != '.docx':
            raise ValueError(f"仅支持.docx格式，当前文件: {self.docx_path.suffix}")
        self.doc = Document(str(self.docx_path))

    def parse(self) -> ParsedDocument:
        """主入口：解析DOCX，返回四部分结构化数据"""
        result = ParsedDocument(file_path=str(self.docx_path))

        # 1. 识别四部分
        section_map = self._identify_sections()

        # 2. 提取每个section的内容
        for stype, sdata in section_map.items():
            self._extract_section_content(sdata)
            result.sections[stype] = sdata

        # 3. 收集所有段落
        result.all_paragraphs = [p.text for p in self.doc.paragraphs]

        # 4. 检测专利类型
        result.patent_type = self._detect_patent_type()
        if result.patent_type == '未知':
            result.warnings.append("未能自动检测专利类型（未发现'本发明'或'本实用新型'）")

        # 5. 提取发明名称
        result.invention_name = self._extract_invention_name()

        # 6. 检查是否有未识别的section
        if not section_map:
            result.warnings.append(
                "未通过页眉识别到任何部分。请确认文件使用了标准分节符和页眉。"
            )

        return result

    # ──────────────────────────────────────────
    # 分节识别
    # ──────────────────────────────────────────

    def _identify_sections(self) -> dict:
        """通过页眉关键词识别四部分"""
        section_map = {}

        # 方法1：通过doc.sections的页眉识别
        doc_sections = self.doc.sections
        for idx, section in enumerate(doc_sections):
            header_text = self._get_header_text(section)
            matched_type = self._match_header(header_text)

            if matched_type:
                sdata = SectionData(
                    section_type=matched_type,
                    header_text=header_text,
                    raw_section=section,
                )
                section_map[matched_type] = sdata

        # 如果通过页眉识别到了四部分，直接返回
        if len(section_map) >= 4:
            self._assign_paragraph_ranges(section_map)
            return section_map

        # 当页眉识别到部分但未识别到摘要时，尝试内容推断摘要
        if len(section_map) >= 2 and '摘要' not in section_map:
            abstract_data = self._detect_abstract_from_content()
            if abstract_data:
                section_map['摘要'] = abstract_data
                # 如果现在有4部分了，直接返回
                if len(section_map) >= 4:
                    self._assign_paragraph_ranges(section_map)
                    return section_map

        # 如果通过页眉+内容推断识别到了≥2部分，继续处理
        if len(section_map) >= 2:
            self._assign_paragraph_ranges(section_map)
            return section_map

        # 方法2：回退到正文标题识别
        header_section_map = dict(section_map)  # 保留页眉识别的结果
        title_sections = self._identify_by_titles()

        for stype, sdata in title_sections.items():
            if stype not in header_section_map:
                header_section_map[stype] = sdata
                header_section_map[stype].header_text = "(正文标题推断)"

        if header_section_map:
            self._assign_paragraph_ranges(header_section_map)
            return header_section_map

        # 方法3：通过内容特征推断
        content_sections = self._identify_by_content(header_section_map)
        if content_sections:
            for stype, sdata in content_sections.items():
                if stype not in header_section_map:
                    header_section_map[stype] = sdata
                    header_section_map[stype].header_text = "(内容特征推断)"
            self._assign_paragraph_ranges(header_section_map)
            return header_section_map

        return section_map

    def _get_header_text(self, section) -> str:
        """获取section的页眉文本"""
        try:
            header = section.header
            if header.is_linked_to_previous:
                return ""
            texts = []
            for p in header.paragraphs:
                if p.text.strip():
                    texts.append(p.text.strip())
            return ' '.join(texts)
        except Exception:
            return ""

    def _match_header(self, header_text: str) -> Optional[str]:
        """匹配页眉文本到section类型"""
        if not header_text:
            return None

        # 先检查"说明书附图"（必须在"说明书"之前检查，因为"说明书附图"也包含"说明书"）
        if self.HEADER_PATTERNS['说明书附图'].search(header_text):
            return '说明书附图'
        if self.HEADER_PATTERNS['摘要'].search(header_text):
            return '摘要'
        if self.HEADER_PATTERNS['权利要求书'].search(header_text):
            return '权利要求书'
        if self.HEADER_PATTERNS['说明书'].search(header_text):
            return '说明书'

        return None

    def _identify_by_titles(self) -> dict:
        """通过正文标题识别四部分"""
        section_map = {}
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            # 检查是否为标题样式（加粗+居中+短文本）
            is_title_like = self._is_title_paragraph(para)
            if not is_title_like:
                continue

            for stype, pattern in self.TITLE_PATTERNS.items():
                if pattern.match(text):
                    section_map[stype] = SectionData(
                        section_type=stype,
                        header_text=f"(正文标题: {text})",
                    )
                    break

        return section_map

    def _detect_abstract_from_content(self) -> Optional[SectionData]:
        """从文档首段内容推断是否为说明书摘要"""
        first_para_text = ""
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if text:
                first_para_text = text
                break

        if not first_para_text:
            return None

        # 摘要特征指示词
        abstract_indicators = [
            '本发明涉及', '本实用新型涉及', '本申请涉及',
            '本发明公开', '本实用新型公开',
        ]
        is_abstract = any(ind in first_para_text for ind in abstract_indicators)
        is_claims = re.match(r'^[1-9][\d]*[\.．、]\s*一\s*种', first_para_text)

        if is_abstract and not is_claims:
            return SectionData(
                section_type='摘要',
                header_text='(内容特征推断: 摘要)',
            )
        return None

    def _identify_by_content(self, existing_map: dict = None) -> dict:
        """通过内容特征推断四部分"""
        section_map = {}
        existing_map = existing_map or {}

        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            # 权利要求编号格式：1.一种... / 1．一种...
            if re.match(r'^[1-9][\d]*[\.．、]\s*一\s*种', text):
                if '权利要求书' not in existing_map and '权利要求书' not in section_map:
                    section_map['权利要求书'] = SectionData(
                        section_type='权利要求书',
                        header_text='(内容特征推断)',
                    )

        return section_map

    def _is_title_paragraph(self, para) -> bool:
        """判断段落是否为标题样式"""
        text = para.text.strip()
        if not text or len(text) > 30:
            return False

        # 检查加粗
        is_bold = False
        for run in para.runs:
            if run.bold:
                is_bold = True
                break

        # 检查居中
        alignment = para.alignment
        is_centered = alignment is not None and str(alignment) == 'CENTER (1)'

        return is_bold or is_centered

    def _assign_paragraph_ranges(self, section_map: dict):
        """为每个section分配段落范围"""
        all_paras = self.doc.paragraphs
        total = len(all_paras)

        # 核心策略：通过页眉切换点 + 正文标题定位来确定段落范围
        # python-docx的section之间有分节符，每个section包含特定段落
        # 我们利用这一点来精确定位

        # 方法1：遍历段落，找到每个section的起始段落
        section_starts = {}

        # 先找页眉切换对应的段落位置
        if len(self.doc.sections) > 0:
            # 收集所有section的页眉和对应的第一个段落位置
            # python-docx中，section的起始位置可以通过分节符在段落中的位置确定
            section_info = []
            for sec_idx, doc_sec in enumerate(self.doc.sections):
                header_text = self._get_header_text(doc_sec)
                matched_type = self._match_header(header_text)
                if matched_type and matched_type in section_map:
                    section_info.append((matched_type, sec_idx, header_text))

            # 通过分节符位置确定段落范围
            # 分节符在某个段落的属性中
            sec_break_paras = []
            for i, para in enumerate(all_paras):
                pPr = para._element.find(qn('w:pPr'))
                if pPr is not None:
                    secPr = pPr.find(qn('w:sectPr'))
                    if secPr is not None:
                        sec_break_paras.append(i)

            # 如果找到了分节符，用它们来划分段落范围
            if sec_break_paras:
                # 第一个section从段落0开始
                boundaries = [0] + [p + 1 for p in sec_break_paras] + [total]
                for idx, (matched_type, sec_idx, header_text) in enumerate(section_info):
                    if idx < len(boundaries) - 1:
                        start = boundaries[sec_idx] if sec_idx < len(boundaries) else 0
                        end = boundaries[sec_idx + 1] - 1 if sec_idx + 1 < len(boundaries) else total - 1
                        section_map[matched_type].para_index_range = (start, end)
                return

        # 方法2：回退到标题定位
        for i, para in enumerate(all_paras):
            text = para.text.strip()
            if not text:
                continue

            for stype, pattern in self.TITLE_PATTERNS.items():
                if pattern.match(text) and stype in section_map:
                    if stype not in section_starts:
                        section_starts[stype] = i

            # 也检查页眉推断的标记
            if '摘要' in section_map and '摘要' not in section_starts:
                if re.match(r'^说?\s*明?\s*书?\s*摘\s*要', text):
                    section_starts['摘要'] = i

        # 按起始位置排序，确定范围
        sorted_starts = sorted(section_starts.items(), key=lambda x: x[1])
        for idx, (stype, start) in enumerate(sorted_starts):
            end = sorted_starts[idx + 1][1] - 1 if idx + 1 < len(sorted_starts) else total - 1
            if stype in section_map:
                section_map[stype].para_index_range = (start, end)

    # ──────────────────────────────────────────
    # 内容提取
    # ──────────────────────────────────────────

    def _extract_section_content(self, sdata: SectionData):
        """从python-docx section对象提取文本/图片/表格/公式"""
        if sdata.raw_section is None:
            # 没有原始section对象，从全局段落中提取
            self._extract_from_paragraphs(sdata)
            return

        section = sdata.raw_section
        paragraphs = []

        # 遍历section中的段落
        # 注意：python-docx的section对象不直接持有段落，需要通过文档全局段落和分节符位置推算
        # 这里改用para_index_range来提取
        self._extract_from_paragraphs(sdata)

    def _extract_from_paragraphs(self, sdata: SectionData):
        """从全局段落中提取指定范围的内容"""
        start, end = sdata.para_index_range
        all_paras = self.doc.paragraphs

        paragraphs = []
        figure_marks = set()
        images = []
        tables = []
        has_omath = False
        img_idx = 0

        for i in range(max(0, start), min(end + 1, len(all_paras))):
            para = all_paras[i]
            text = para.text.strip()
            paragraphs.append(text)

            # 提取附图标记
            marks = self._extract_figure_marks(text)
            figure_marks.update(marks)

            # 检测OOXML公式
            if self._has_omath(para):
                has_omath = True

            # 提取图片
            for run in para.runs:
                for drawing in run._element.findall(qn('w:drawing')):
                    img_info = self._extract_image_from_drawing(drawing, img_idx)
                    if img_info:
                        images.append(img_info)
                        img_idx += 1

                # 也检查inline shapes
                for inline in run._element.findall(qn('w:pict')):
                    img_info = self._extract_image_from_pict(inline, img_idx)
                    if img_info:
                        images.append(img_info)
                        img_idx += 1

        sdata.paragraphs = paragraphs
        sdata.figure_marks = sorted(figure_marks, key=lambda x: int(x) if x.isdigit() else x)
        sdata.images = images
        sdata.has_omath = has_omath

        # 提取表格（如果在段落范围内有表格）
        for table in self.doc.tables:
            table_text = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_text.append(row_data)
            if table_text:
                # 检查表格是否在当前section范围内（启发式：检查第一行文本是否在段落中）
                first_cell = table_text[0][0] if table_text else ""
                if any(first_cell[:20] in p for p in paragraphs if p):
                    tables.append(table_text)

        sdata.tables = tables

    def _extract_figure_marks(self, text: str) -> list:
        """从文本中提取附图标记

        提取策略：
        1. 括号内数字 (1)、（2）— 最明确的附图标记格式
        2. "图N"格式 — 附图引用格式
        3. 【已禁用】中文后的2-4位数字 — 该正则（FIGURE_MARK_IN_TEXT）
           对软件/算法类专利误报率极高（如"20人""50次""15毫秒"等普通数值
           被误判为附图标记），故不再使用。机械/结构类专利的组件编号
           通常也使用括号格式，不影响F-83核查。
        """
        marks = set()

        # 1. 括号内数字： (101)、（201）
        for m in self.FIGURE_MARK_IN_BRACKETS.finditer(text):
            val = m.group(1) or m.group(2)
            if val and len(val) <= 4:
                marks.add(val)

        # 2. "图N"格式
        for m in re.finditer(r'图\s*(\d{1,3})', text):
            marks.add(m.group(1))

        return list(marks)

    def _has_omath(self, paragraph) -> bool:
        """检查段落是否含OOXML公式"""
        try:
            omaths = paragraph._element.findall(
                './/{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath'
            )
            return len(omaths) > 0
        except Exception:
            return False

    def _extract_image_from_drawing(self, drawing, index: int) -> Optional[ImageInfo]:
        """从w:drawing元素提取图片"""
        try:
            blip = drawing.find('.//' + qn('a:blip'))
            if blip is None:
                blip = drawing.find(
                    './/{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
                )
            if blip is None:
                return None

            rId = blip.get(qn('r:embed')) or blip.get(
                '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
            )
            if not rId:
                return None

            # 获取图片数据
            image_part = self.doc.part.related_parts[rId]
            img_data = image_part.blob
            content_type = image_part.content_type

            # 确定格式
            fmt = 'unknown'
            if 'png' in content_type:
                fmt = 'png'
            elif 'jpeg' in content_type or 'jpg' in content_type:
                fmt = 'jpeg'
            elif 'emf' in content_type:
                fmt = 'emf'
            elif 'wmf' in content_type:
                fmt = 'wmf'
            elif 'tiff' in content_type:
                fmt = 'tiff'
            elif 'bmp' in content_type:
                fmt = 'bmp'

            # 获取尺寸
            extent = drawing.find('.//' + qn('wp:extent'))
            width_pt = 0.0
            height_pt = 0.0
            if extent is not None:
                cx = extent.get('cx')
                cy = extent.get('cy')
                if cx:
                    width_pt = int(cx) / 12700  # EMU → pt (1pt = 12700 EMU)
                if cy:
                    height_pt = int(cy) / 12700

            return ImageInfo(
                index=index,
                data=img_data,
                width_pt=width_pt,
                height_pt=height_pt,
                format=fmt,
            )
        except Exception:
            return None

    def _extract_image_from_pict(self, pict, index: int) -> Optional[ImageInfo]:
        """从w:pict元素提取图片"""
        try:
            imagedata = pict.find('.//' + qn('v:imagedata'))
            if imagedata is None:
                return None

            rId = imagedata.get(qn('r:id'))
            if not rId:
                return None

            image_part = self.doc.part.related_parts[rId]
            img_data = image_part.blob
            content_type = image_part.content_type

            fmt = 'unknown'
            if 'png' in content_type:
                fmt = 'png'
            elif 'jpeg' in content_type or 'jpg' in content_type:
                fmt = 'jpeg'
            elif 'emf' in content_type:
                fmt = 'emf'
            elif 'wmf' in content_type:
                fmt = 'wmf'

            return ImageInfo(
                index=index,
                data=img_data,
                width_pt=0,
                height_pt=0,
                format=fmt,
            )
        except Exception:
            return None

    # ──────────────────────────────────────────
    # 专利类型检测
    # ──────────────────────────────────────────

    def _detect_patent_type(self) -> str:
        """从内容自动检测专利类型"""
        count_utility = 0  # "本实用新型"
        count_invention = 0  # "本发明"

        for para in self.doc.paragraphs:
            text = para.text
            count_utility += len(self.PATENT_TYPE_PATTERNS['实用新型'].findall(text))
            count_invention += len(self.PATENT_TYPE_PATTERNS['发明'].findall(text))

        # 排除"本发明"在"本实用新型"中的子匹配
        # "本实用新型"会被两个pattern都匹配，需要调整
        # 实际上"本实用新型"不会匹配"本发明"的pattern，因为中间有"实用"二字
        # 但"本发明内容"等短语中的"本发明"会被匹配
        # 策略：统计纯"本发明"（不含"本实用新型"）的出现次数

        if count_utility > 0 and count_invention == 0:
            return '实用新型'
        elif count_invention > 0 and count_utility == 0:
            return '发明'
        elif count_utility > 0 and count_invention > 0:
            # 两者都有 → 类型用语矛盾，标记但返回出现更多的
            if count_utility > count_invention:
                return '实用新型'
            else:
                return '发明'
        else:
            return '未知'

    def _extract_invention_name(self) -> str:
        """提取发明名称（优先从说明书部分的首段）"""
        # 策略1：优先从说明书section的第一段提取
        # 先通过页眉识别说明书section的段落范围
        spec_start = None
        spec_end = None
        for idx, section in enumerate(self.doc.sections):
            header_text = self._get_header_text(section)
            if self.HEADER_PATTERNS['说明书'].search(header_text) and \
               not self.HEADER_PATTERNS['说明书附图'].search(header_text):
                # 找到了说明书section，计算其段落范围
                sec_break_paras = []
                for i, para in enumerate(self.doc.paragraphs):
                    pPr = para._element.find(qn('w:pPr'))
                    if pPr is not None:
                        secPr = pPr.find(qn('w:sectPr'))
                        if secPr is not None:
                            sec_break_paras.append(i)
                boundaries = [0] + [p + 1 for p in sec_break_paras] + [len(self.doc.paragraphs)]
                if idx < len(boundaries) - 1:
                    spec_start = boundaries[idx] if idx < len(boundaries) else 0
                    spec_end = boundaries[idx + 1] - 1 if idx + 1 < len(boundaries) else len(self.doc.paragraphs) - 1
                break

        # 排除已知的标题词
        skip_patterns = [
            r'^说\s*明\s*书\s*摘\s*要$',
            r'^权\s*利\s*要\s*求\s*书$',
            r'^说\s*明\s*书$',
            r'^说\s*明\s*书\s*附\s*图$',
            r'^技\s*术\s*领\s*域$',
            r'^背\s*景\s*技\s*术$',
            r'^(发\s*明|实\s*用\s*新\s*型)\s*内\s*容$',
            r'^附\s*图\s*说\s*明$',
            r'^具\s*体\s*实\s*施\s*方\s*式$',
        ]

        def _is_invention_name(text: str) -> bool:
            if not text or len(text) < 5 or len(text) > 60:
                return False
            for sp in skip_patterns:
                if re.match(sp, text):
                    return False
            return ('一种' in text or '方法' in text or
                    '装置' in text or '系统' in text or
                    '设备' in text or '产品' in text)

        # 策略1：从说明书section查找
        if spec_start is not None and spec_end is not None:
            for i in range(spec_start, min(spec_end + 1, len(self.doc.paragraphs))):
                text = self.doc.paragraphs[i].text.strip()
                if _is_invention_name(text):
                    return text

        # 策略2：从全文查找（回退）
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if _is_invention_name(text):
                return text

        return ""

    # ──────────────────────────────────────────
    # 工具方法
    # ──────────────────────────────────────────

    def get_section_summary(self, result: ParsedDocument) -> str:
        """生成分节确认摘要"""
        lines = []
        type_emoji = {'摘要': '📋', '权利要求书': '⚖️', '说明书': '📖', '说明书附图': '🖼️'}

        for stype in ['摘要', '权利要求书', '说明书', '说明书附图']:
            if stype in result.sections:
                sec = result.sections[stype]
                emoji = type_emoji.get(stype, '📄')
                start, end = sec.para_index_range
                lines.append(
                    f"✅ {emoji} [{stype}] 页眉=\"{sec.header_text}\" "
                    f"（第{start+1}-{end+1}段）"
                )

        lines.append(f"📋 专利类型检测：{result.patent_type}")
        if result.invention_name:
            lines.append(f"📋 发明名称：{result.invention_name}")

        for w in result.warnings:
            lines.append(f"⚠️ {w}")

        return '\n'.join(lines)


# ──────────────────────────────────────────────
# 命令行入口
# ──────────────────────────────────────────────

def main():
    """命令行测试入口"""
    if len(sys.argv) < 2:
        print("用法: python docx_parser.py <docx文件路径>")
        sys.exit(1)

    path = sys.argv[1]
    parser = DocxSectionParser(path)
    result = parser.parse()

    # 输出摘要
    print(parser.get_section_summary(result))

    # 输出详细信息
    for stype, sec in result.sections.items():
        print(f"\n{'='*60}")
        print(f"部分：{stype}")
        print(f"段落数：{len(sec.paragraphs)}")
        print(f"附图标记：{sec.figure_marks}")
        print(f"图片数量：{len(sec.images)}")
        print(f"含OOXML公式：{sec.has_omath}")
        print(f"表格数量：{len(sec.tables)}")
        # 打印前5段
        for i, p in enumerate(sec.paragraphs[:5]):
            if p:
                print(f"  段落{i}: {p[:80]}")


if __name__ == '__main__':
    main()
